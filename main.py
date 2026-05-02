import asyncio
import io
import hashlib
import json
import logging
import os
import mimetypes
import re
import uuid
import urllib.error
import urllib.request
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Optional

import openai
import openpyxl
import pdfplumber
import docx
from dotenv import load_dotenv

from fastapi import FastAPI, HTTPException, UploadFile, File, Request, Form
from fastapi.responses import HTMLResponse, JSONResponse, Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from parser_edital import (
    analisar_sem_api,
    gerar_ficha,
    _is_identificado,
    calcular_confianca,
    calcular_score_viabilidade,
)

try:
    import fitz
except Exception:
    fitz = None

try:
    import numpy as np
except Exception:
    np = None

try:
    from rapidocr_onnxruntime import RapidOCR
except Exception:
    RapidOCR = None

load_dotenv()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(name)s [%(request_id)s] %(message)s",
)


class _RequestIdFilter(logging.Filter):
    def filter(self, record):
        if not hasattr(record, "request_id"):
            record.request_id = "-"
        return True


for _handler in logging.getLogger().handlers:
    _handler.addFilter(_RequestIdFilter())

logger = logging.getLogger("licitapro")


def _audit(event: str, request_id: str = "-", **fields):
    payload = {"event": event, "request_id": request_id}
    for key, value in fields.items():
        if value is not None:
            payload[key] = value
    logger.info(json.dumps(payload, ensure_ascii=False, separators=(",", ":")), extra={"request_id": request_id})

app = FastAPI(title="Proxy LicitaÃ§Ã£o")

import datetime

LIMITE_DIARIO = 20  # anÃ¡lises por dia
USAR_PARSER_LOCAL = os.getenv("USAR_PARSER_LOCAL", "true").lower() not in ("0", "false", "no", "off")
PARSER_FALLBACK_API = os.getenv("PARSER_FALLBACK_API", "true").lower() not in ("0", "false", "no", "off")
try:
    PARSER_MIN_CONFIANCA = int(os.getenv("PARSER_MIN_CONFIANCA", "70"))
except ValueError:
    PARSER_MIN_CONFIANCA = 70
try:
    PARSER_MAX_CHARS_FALLBACK = int(os.getenv("PARSER_MAX_CHARS_FALLBACK", "80000"))
except ValueError:
    PARSER_MAX_CHARS_FALLBACK = 80000
OCR_HABILITADO = os.getenv("OCR_HABILITADO", "true").lower() not in ("0", "false", "no", "off")
try:
    OCR_MIN_CHAR = int(os.getenv("OCR_MIN_CHAR", "120"))
except ValueError:
    OCR_MIN_CHAR = 120
try:
    OCR_MAX_PAGINAS = int(os.getenv("OCR_MAX_PAGINAS", "20"))
except ValueError:
    OCR_MAX_PAGINAS = 20
try:
    OCR_DPI = int(os.getenv("OCR_DPI", "220"))
except ValueError:
    OCR_DPI = 220
try:
    OCR_MAX_FILE_BYTES = int(os.getenv("OCR_MAX_FILE_BYTES", "819200"))  # 800 KB
except ValueError:
    OCR_MAX_FILE_BYTES = 819200
try:
    MAX_PAGINAS_GRANDES = int(os.getenv("MAX_PAGINAS_GRANDES", "15"))
except ValueError:
    MAX_PAGINAS_GRANDES = 15
APP_VERSION = os.getenv("APP_VERSION", "dev")
APP_CHANNEL = os.getenv("APP_CHANNEL", "local")
APP_COMMIT = os.getenv("APP_COMMIT") or os.getenv("RENDER_GIT_COMMIT") or "local"
APP_DEPLOYED_AT = os.getenv("APP_DEPLOYED_AT") or os.getenv("RENDER_DEPLOYED_AT") or ""
APP_VERSION_LABEL = f"{APP_VERSION} Â· {APP_CHANNEL}"
APP_COMMIT_LABEL = APP_COMMIT[:7] if APP_COMMIT else "local"

_OCR_ENGINE = None

# Fallback local (LLM rodando localmente, ex: Ollama ou endpoint OpenAI-compat).
LOCAL_LLM_ENABLED = os.getenv("LOCAL_LLM_ENABLED", "false").lower() in ("1", "true", "yes", "on")
LOCAL_LLM_PROVIDER = os.getenv("LOCAL_LLM_PROVIDER", "ollama").lower()  # "ollama" | "openai_compat"
LOCAL_LLM_BASE_URL = os.getenv("LOCAL_LLM_BASE_URL", "http://127.0.0.1:11434")
LOCAL_LLM_MODEL = os.getenv("LOCAL_LLM_MODEL", "llama3.1:8b")
LOCAL_LLM_TIMEOUT_S = float(os.getenv("LOCAL_LLM_TIMEOUT_S", "30"))

_stats = {
    "total_analises": 0,
    "tokens_input_total": 0,
    "tokens_output_total": 0,
    "custo_usd_total": 0.0,
    "por_provedor": {},
    "hoje": datetime.date.today().isoformat(),
    "analises_hoje": 0,
}

# custo em USD por 1M tokens (input, output)
_CUSTO = {
    "gpt-4.1-nano":                           (0.10,  0.40),  # mais barato
    "gpt-4o-mini":                            (0.15,  0.60),
    "google/gemma-3-27b-it:free":             (0.0,   0.0),
    "meta-llama/llama-3.3-70b-instruct:free": (0.0,   0.0),
    "google/gemma-4-26b-a4b-it:free":         (0.0,   0.0),
    "nvidia/nemotron-3-super-120b-a12b:free": (0.0,   0.0),
    "google/gemma-3-12b-it:free":             (0.0,   0.0),
    "llama-3.3-70b-versatile":                (0.0,   0.0),
}

_openai_api_key = os.getenv("OPENAI_API_KEY") or ""
_openai = openai.AsyncOpenAI(api_key=_openai_api_key) if _openai_api_key.strip() else None
_gemini_api_key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
_openrouter_api_key = os.getenv("OPENROUTER_API_KEY") or ""
_openrouter = (
    openai.AsyncOpenAI(api_key=_openrouter_api_key, base_url="https://openrouter.ai/api/v1")
    if _openrouter_api_key.strip()
    else None
)
_groq_api_key = os.getenv("GROQ_API_KEY") or ""
_groq = (
    openai.AsyncOpenAI(api_key=_groq_api_key, base_url="https://api.groq.com/openai/v1")
    if _groq_api_key.strip()
    else None
)
_groq2_api_key = os.getenv("GROQ_API_KEY2") or ""
_groq2 = (
    openai.AsyncOpenAI(api_key=_groq2_api_key, base_url="https://api.groq.com/openai/v1")
    if _groq2_api_key.strip()
    else None
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.middleware("http")
async def request_logging_middleware(request: Request, call_next):
    request_id = uuid.uuid4().hex[:8]
    request.state.request_id = request_id
    try:
        response = await call_next(request)
    except Exception:
        logger.exception("Unhandled request failure", extra={"request_id": request_id})
        return JSONResponse(
            status_code=500,
            content={"detail": "Erro interno ao processar a requisiÃ§Ã£o.", "request_id": request_id},
            headers={"X-Request-ID": request_id},
        )
    response.headers["X-Request-ID"] = request_id
    return response

SYSTEM_PROMPT = """VocÃª Ã© um especialista em licitaÃ§Ãµes pÃºblicas brasileiras.
REGRA ABSOLUTA: comece a resposta IMEDIATAMENTE com "## FICHA DE LICITAÃ‡ÃƒO" â€” ZERO introduÃ§Ãµes, ZERO raciocÃ­nio, ZERO explicaÃ§Ãµes. Apenas a ficha em Markdown.
Se um dado realmente nÃ£o constar em nenhuma parte do texto, escreva "NÃ£o informado". Mas ANTES de escrever "NÃ£o informado" em qualquer campo, leia o documento INTEIRO â€” especialmente seÃ§Ãµes de habilitaÃ§Ã£o, anexos, clÃ¡usulas e condiÃ§Ãµes. Ã‰ PROIBIDO escrever "NÃ£o informado" para Documentos de HabilitaÃ§Ã£o se o edital contiver qualquer lista de documentos exigidos.
Calcule o Valor Total de cada item (Qtd Ã— Valor Unit.).
Siga EXATAMENTE esta estrutura:

## FICHA DE LICITAÃ‡ÃƒO

| Campo | Valor |
|---|---|
| **NÂº / Processo** | ... |
| **Ã“rgÃ£o** | ... |
| **Modalidade** | ... |
| **CritÃ©rio de Julgamento** | ... |
| **Valor Estimado Total** | R$ ... |
| **VigÃªncia do Contrato** | ... |
| **Abertura das Propostas** | ... |
| **Prazo para Envio de Proposta** | ... |

## Objeto
[descriÃ§Ã£o completa do objeto licitado]

## CondiÃ§Ãµes Financeiras
- **Garantia Contratual:** ...
- **Prazo de Pagamento:** ...
- **PatrimÃ´nio LÃ­quido MÃ­nimo:** ...
- **Capital Social MÃ­nimo:** ...

## Posto de Atendimento
[local/endereÃ§o onde os serviÃ§os serÃ£o prestados ou as entregas realizadas]

## Contato do Ã“rgÃ£o
[e-mail e telefone do responsÃ¡vel]

## Itens a Cotar

| # | DescriÃ§Ã£o | Unid. | Qtd. | Valor Unit. | Valor Total |
|---|-----------|-------|------|-------------|-------------|
| 1 | ... | ... | ... | R$ ... | R$ ... |

## Modelo de Proposta
[tipo de taxa, o que Ã© cotado, como lanÃ§ar, observaÃ§Ãµes importantes]

## Documentos de HabilitaÃ§Ã£o

### JurÃ­dica
- ...

### Fiscal e Trabalhista
- ...

### EconÃ´mico-Financeira
- ...

### TÃ©cnica
- ...

### Outras ExigÃªncias
- ...

## âš ï¸ Alertas
> [ponto crÃ­tico relevante para a decisÃ£o de participar]

> [prÃ³ximo alerta, se houver]

## Score de Viabilidade
**Score:** [inteiro de 0 a 100]
**NÃ­vel:** [Alta | MÃ©dia | Baixa]
**Justificativa:** [2 linhas diretas avaliando acessibilidade do objeto, dificuldade das exigÃªncias e oportunidade geral]

## AnÃ¡lise de ExigÃªncias
[Liste cada exigÃªncia de habilitaÃ§Ã£o com prefixo de status:]
[ok] exigÃªncia â€” orientaÃ§Ã£o breve
[warn] exigÃªncia â€” por que merece atenÃ§Ã£o
[fail] exigÃªncia â€” por que Ã© restritiva

CritÃ©rios: [ok]=padrÃ£o fÃ¡cil (certidÃµes online, CNPJ, atestados genÃ©ricos); [warn]=requer preparaÃ§Ã£o, prazo <15 dias, PL especÃ­fico; [fail]=muito restritivo, volume >5000 unidades, PL >R$500k, sede especÃ­fica, requisito raro."""

USER_TEMPLATE = "Analise o seguinte edital ({num_docs} documento(s)) e gere a ficha:\n\n{texto}"

HTML_PAGE = """<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>LicitaPro â€” AnÃ¡lise de Editais</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap" rel="stylesheet">
<script src="https://cdn.jsdelivr.net/npm/marked/marked.min.js"></script>
<script src="https://cdn.jsdelivr.net/npm/pdfjs-dist@3.11.174/build/pdf.min.js"></script>
<style>
:root{
  --brand-900:#061A33;--brand-800:#0A2540;--brand-700:#0E335A;
  --brand-600:#1E4FA0;--brand-500:#2F6FE0;--brand-400:#5C90EE;
  --brand-300:#8FB3F4;--brand-200:#C5D7F9;--brand-100:#E6EEFC;--brand-50:#F4F8FE;
  --ink-900:#111827;--ink-800:#1F2937;--ink-700:#374151;--ink-600:#4B5563;
  --ink-500:#6B7280;--ink-400:#9CA3AF;--ink-300:#D1D5DB;--ink-200:#E5E7EB;
  --ink-150:#EEF0F3;--ink-100:#F3F4F6;--ink-50:#F9FAFB;--ink-0:#FFFFFF;
  --success-700:#166534;--success-500:#22C55E;--success-100:#DCFCE7;--success-50:#F0FDF4;
  --warn-700:#92400E;--warn-500:#F59E0B;--warn-100:#FEF3C7;--warn-50:#FFFBEB;
  --danger-700:#991B1B;--danger-500:#EF4444;--danger-100:#FEE2E2;--danger-50:#FEF2F2;
  --bg:#FFFFFF;--bg-subtle:#FAFBFC;
  --border:#E5E7EB;--border-subtle:#EEF0F3;
  --fg-1:#111827;--fg-2:#4B5563;--fg-3:#6B7280;--fg-4:#9CA3AF;
  --font-sans:'Inter',-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;
  --font-mono:'JetBrains Mono',Menlo,Consolas,monospace;
  --radius-sm:6px;--radius-md:10px;--radius-lg:14px;--radius-xl:20px;--radius-full:9999px;
  --shadow-sm:0 1px 2px rgba(11,15,20,.04),0 1px 3px rgba(11,15,20,.06);
  --shadow-md:0 4px 8px -2px rgba(11,15,20,.06),0 2px 4px -2px rgba(11,15,20,.04);
  --shadow-lg:0 12px 24px -8px rgba(11,15,20,.10),0 4px 8px -4px rgba(11,15,20,.06);
  --ease-out:cubic-bezier(0.22,1,0.36,1);
  --sidebar-w:220px;
}
*{margin:0;padding:0;box-sizing:border-box}
body{font-family:var(--font-sans);background:var(--bg-subtle);color:var(--fg-1);min-height:100vh;-webkit-font-smoothing:antialiased;overflow-x:hidden}
.app{display:flex;min-height:100vh}
.sidebar{width:var(--sidebar-w);flex-shrink:0;background:var(--bg);border-right:1px solid var(--border-subtle);display:flex;flex-direction:column;position:fixed;top:0;left:0;height:100vh;z-index:50}
.sidebar-logo{padding:20px 20px 16px;font-size:17px;font-weight:700;color:var(--ink-900);letter-spacing:-.025em;border-bottom:1px solid var(--border-subtle);user-select:none}
.sidebar-logo span{color:var(--brand-500)}
.sidebar-nav{flex:1;padding:12px 10px;display:flex;flex-direction:column;gap:2px}
.nav-item{display:flex;align-items:center;gap:10px;padding:8px 10px;border-radius:8px;background:transparent;border:none;color:var(--fg-2);font-size:13px;font-weight:500;font-family:var(--font-sans);cursor:pointer;text-align:left;text-decoration:none;transition:background 120ms var(--ease-out),color 120ms var(--ease-out)}
.nav-item:hover{background:var(--ink-50);color:var(--fg-1)}
.nav-item.active{background:var(--brand-50);color:var(--brand-700);font-weight:600}
.nav-item svg{flex-shrink:0;width:16px;height:16px}
.sidebar-footer{padding:12px 14px 16px;border-top:1px solid var(--border-subtle)}
.quota-label{font-size:11px;font-weight:600;letter-spacing:.06em;text-transform:uppercase;color:var(--fg-3);margin-bottom:6px}
.quota-val{font-size:13px;font-weight:600;color:var(--fg-1);font-variant-numeric:tabular-nums}
.quota-bar{height:4px;background:var(--ink-150);border-radius:9999px;margin-top:6px;overflow:hidden}
.quota-fill{height:100%;border-radius:9999px;background:var(--brand-500);transition:width .3s}
.version-line{margin-top:10px;font-size:11px;font-weight:600;letter-spacing:.06em;text-transform:uppercase;color:var(--fg-3)}
.version-sub{margin-top:4px;font-size:11px;color:var(--fg-4);font-variant-numeric:tabular-nums;white-space:nowrap;overflow:hidden;text-overflow:ellipsis}
.main-content{margin-left:var(--sidebar-w);flex:1;min-height:100vh;overflow-y:auto}
.page{padding:40px 48px;max-width:1100px;animation:fadeIn 200ms var(--ease-out)}
@keyframes fadeIn{from{opacity:0;transform:translateY(6px)}to{opacity:1;transform:none}}
.page-eyebrow{font-size:11px;color:var(--brand-500);font-weight:600;letter-spacing:.12em;text-transform:uppercase;margin-bottom:8px}
.page-title{font-size:32px;font-weight:700;letter-spacing:-.025em;color:var(--fg-1);line-height:1.15}
.page-sub{margin-top:6px;font-size:15px;color:var(--fg-2)}
.page-header{display:flex;align-items:flex-end;justify-content:space-between;gap:24px;margin-bottom:32px}
.stats-grid{display:grid;gap:14px;margin-bottom:28px}
.g4{grid-template-columns:repeat(4,1fr)}
.g3{grid-template-columns:repeat(3,1fr)}
.g2{grid-template-columns:1fr 1fr}
.stat-card{background:var(--bg);border:1px solid var(--border);border-radius:var(--radius-lg);padding:18px 20px;box-shadow:var(--shadow-sm)}
.stat-card .lbl{font-size:11px;font-weight:600;letter-spacing:.06em;text-transform:uppercase;color:var(--fg-3);margin-bottom:8px}
.stat-card .val{font-size:26px;font-weight:700;letter-spacing:-.02em;color:var(--fg-1);line-height:1;font-variant-numeric:tabular-nums}
.stat-card .sub{font-size:12px;color:var(--fg-4);margin-top:5px;font-variant-numeric:tabular-nums}
.filter-row{display:flex;align-items:center;gap:8px;margin-bottom:16px;flex-wrap:wrap}
.filter-pill{background:#fff;color:var(--fg-2);border:1px solid var(--border);border-radius:9999px;padding:6px 14px;font-family:var(--font-sans);font-size:13px;font-weight:500;cursor:pointer;transition:all 120ms var(--ease-out)}
.filter-pill.active{background:var(--ink-900);color:#fff;border-color:var(--ink-900)}
.filter-right{margin-left:auto;display:flex;gap:8px}
.edital-list{display:flex;flex-direction:column;gap:12px}
.edital-card{background:#fff;border:1px solid var(--border);border-radius:var(--radius-lg);padding:20px;cursor:pointer;box-shadow:var(--shadow-sm);display:grid;grid-template-columns:72px 1fr auto;gap:20px;align-items:center;transition:box-shadow 200ms var(--ease-out),border-color 200ms var(--ease-out)}
.edital-card:hover{box-shadow:var(--shadow-md);border-color:var(--ink-300)}
.edital-meta{display:flex;align-items:center;gap:8px;margin-bottom:6px;flex-wrap:wrap}
.edital-numero{font-family:var(--font-mono);font-size:12px;color:var(--fg-3)}
.edital-objeto{font-size:16px;font-weight:600;color:var(--fg-1);line-height:1.3;margin-bottom:4px}
.edital-orgao{font-size:13px;color:var(--fg-2)}
.edital-right{text-align:right;min-width:140px}
.edital-valor{font-size:18px;font-weight:700;font-variant-numeric:tabular-nums;letter-spacing:-.01em;color:var(--fg-1)}
.badge{display:inline-flex;align-items:center;padding:2px 10px;border-radius:var(--radius-full);font-size:11px;font-weight:600;white-space:nowrap}
.badge-brand{background:var(--brand-50);color:var(--brand-700)}
.badge-success{background:var(--success-100);color:var(--success-700)}
.badge-warn{background:var(--warn-100);color:var(--warn-700)}
.badge-danger{background:var(--danger-100);color:var(--danger-700)}
.badge-solid{background:var(--ink-900);color:#fff;letter-spacing:.06em}
.btn{display:inline-flex;align-items:center;gap:8px;padding:9px 18px;border-radius:var(--radius-md);font-size:14px;font-weight:600;font-family:var(--font-sans);cursor:pointer;transition:all 120ms var(--ease-out);border:none;text-decoration:none}
.btn-primary{background:var(--brand-500);color:#fff}
.btn-primary:hover{background:var(--brand-600)}
.btn-secondary{background:#fff;color:var(--fg-2);border:1px solid var(--border)}
.btn-secondary:hover{background:var(--ink-50);color:var(--fg-1)}
.btn-sm{padding:7px 13px;font-size:13px}
.btn:active{transform:scale(.98)}
.dropzone{border:2px dashed var(--brand-300);border-radius:var(--radius-xl);padding:64px 32px;background:var(--brand-50);text-align:center;cursor:pointer;transition:all 200ms var(--ease-out)}
.dropzone.over,.dropzone:hover{background:var(--brand-100);border-color:var(--brand-500)}
.dz-icon{width:56px;height:56px;border-radius:9999px;background:#fff;border:1px solid var(--border);display:inline-flex;align-items:center;justify-content:center;color:var(--brand-500);margin:0 auto 20px}
.dz-title{font-size:20px;font-weight:600;letter-spacing:-.015em;color:var(--fg-1);margin-bottom:8px}
.dz-sub{font-size:14px;color:var(--fg-2)}
.file-list{margin-top:14px;display:flex;flex-direction:column;gap:8px}
.file-item{display:flex;align-items:center;gap:10px;padding:10px 14px;background:var(--brand-50);border:1px solid var(--brand-100);border-radius:var(--radius-sm);font-size:13px}
.file-name{flex:1;overflow:hidden;text-overflow:ellipsis;white-space:nowrap;color:var(--fg-1)}
.file-size{color:var(--fg-4);font-family:var(--font-mono);font-size:11px;white-space:nowrap}
.file-rm{cursor:pointer;color:var(--fg-4);background:none;border:none;padding:2px;border-radius:4px;line-height:0;transition:color 120ms}
.file-rm:hover{color:#EF4444}
.exig-item{display:flex;gap:12px;padding:12px 14px;border-radius:10px;margin-bottom:8px}
.exig-ok{background:var(--success-50);border:1px solid #BBF7D0}
.exig-warn{background:var(--warn-50);border:1px solid #FDE68A}
.exig-fail{background:var(--danger-50);border:1px solid #FECACA}
.exig-icon{flex-shrink:0;margin-top:2px}
.exig-title{font-size:14px;font-weight:600;margin-bottom:2px}
.exig-ok .exig-title{color:var(--success-700)}
.exig-warn .exig-title{color:var(--warn-700)}
.exig-fail .exig-title{color:var(--danger-700)}
.exig-detail{font-size:13px;line-height:1.5;opacity:.85}
.exig-ok .exig-detail{color:var(--success-700)}
.exig-warn .exig-detail{color:var(--warn-700)}
.exig-fail .exig-detail{color:var(--danger-700)}
.back-btn{display:inline-flex;align-items:center;gap:6px;background:transparent;border:none;color:var(--fg-2);font-size:13px;font-family:var(--font-sans);cursor:pointer;padding:0;margin-bottom:24px;transition:color 120ms}
.back-btn:hover{color:var(--fg-1)}
.ficha-content{background:#fff;border:1px solid var(--border);border-radius:var(--radius-lg);padding:28px 32px;box-shadow:var(--shadow-sm)}
.ficha-content h2{font-size:16px;font-weight:700;color:var(--fg-1);margin:24px 0 10px;letter-spacing:-.015em}
.ficha-content h2:first-child{margin-top:0}
.ficha-content h3{font-size:14px;font-weight:600;color:var(--fg-2);margin:14px 0 6px}
.ficha-content p{font-size:14px;line-height:1.6;color:var(--fg-2);margin-bottom:8px}
.ficha-content table{width:100%;border-collapse:collapse;font-size:13px;margin-bottom:12px}
.ficha-content th{padding:9px 12px;text-align:left;border-bottom:2px solid var(--border);font-weight:600;font-size:11px;text-transform:uppercase;letter-spacing:.06em;color:var(--fg-3);background:var(--ink-50)}
.ficha-content td{padding:9px 12px;border-bottom:1px solid var(--border-subtle);color:var(--fg-1);vertical-align:top;font-size:13px}
.ficha-content tr:last-child td{border-bottom:none}
.ficha-content ul,.ficha-content ol{padding-left:20px;margin-bottom:8px}
.ficha-content li{font-size:13px;color:var(--fg-2);line-height:1.6;margin-bottom:4px}
.ficha-content blockquote{border-left:3px solid var(--warn-500);background:var(--warn-50);padding:10px 14px;border-radius:0 6px 6px 0;margin:8px 0}
.ficha-content blockquote p{color:var(--warn-700);margin:0;font-size:13px}
.ficha-content strong{color:var(--fg-1)}
.ficha-content code{font-family:var(--font-mono);font-size:12px;background:var(--ink-100);padding:1px 5px;border-radius:4px}
.processing-card{background:#fff;border:1px solid var(--border);border-radius:var(--radius-xl);padding:64px 48px;text-align:center;box-shadow:var(--shadow-sm)}
.processing-icon{width:56px;height:56px;border-radius:9999px;background:var(--brand-50);display:inline-flex;align-items:center;justify-content:center;color:var(--brand-500);margin:0 auto 20px}
.processing-bar{height:4px;background:var(--ink-150);border-radius:9999px;overflow:hidden;max-width:280px;margin:20px auto 0}
.processing-fill{height:100%;background:var(--brand-500);border-radius:9999px;animation:prog 2s ease-in-out infinite alternate}
@keyframes prog{from{width:20%}to{width:90%}}
#toast{position:fixed;bottom:24px;right:24px;z-index:1000;display:flex;flex-direction:column;gap:8px}
.toast-item{background:var(--ink-900);color:#fff;padding:12px 16px;border-radius:var(--radius-md);font-size:13px;font-weight:500;box-shadow:var(--shadow-lg);animation:toastIn 200ms var(--ease-out)}
@keyframes toastIn{from{transform:translateX(100%);opacity:0}to{transform:none;opacity:1}}
.hist-table{background:#fff;border:1px solid var(--border);border-radius:var(--radius-lg);overflow:hidden;box-shadow:var(--shadow-sm)}
.hist-table table{width:100%;border-collapse:collapse}
.hist-table th{padding:12px 20px;text-align:left;border-bottom:1px solid var(--border);background:var(--ink-50);font-size:11px;font-weight:600;text-transform:uppercase;letter-spacing:.06em;color:var(--fg-3)}
.hist-table td{padding:13px 20px;border-bottom:1px solid var(--border-subtle);font-size:14px;font-variant-numeric:tabular-nums;color:var(--fg-1)}
.hist-table tr:last-child td{border-bottom:none}
.hist-table tr:hover td{background:var(--ink-50)}
.empty-state{text-align:center;padding:64px 32px;color:var(--fg-3)}
.empty-title{font-size:16px;font-weight:600;color:var(--fg-2);margin-top:16px;margin-bottom:6px}
.empty-sub{font-size:14px}
.action-row{display:flex;gap:12px;justify-content:flex-end;margin-top:32px;padding-top:24px;border-top:1px solid var(--border-subtle)}
@media(max-width:900px){
  .page{padding:24px 20px}
  .g4{grid-template-columns:1fr 1fr}
  .edital-card{grid-template-columns:56px 1fr}
  .edital-right{display:none}
}
@media(max-width:640px){
  .sidebar{display:none}
  .main-content{margin-left:0}
  .g4,.g3{grid-template-columns:1fr 1fr}
  .page{padding:20px 16px}
}
</style>
</head>
<body>
<div class="app">

<aside class="sidebar">
  <div class="sidebar-logo">Licita<span>Pro</span></div>
  <nav class="sidebar-nav">
    <button class="nav-item" id="nav-editais" onclick="showPage('editais')">
      <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.75" stroke-linecap="round" stroke-linejoin="round"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/><line x1="16" y1="13" x2="8" y2="13"/><line x1="16" y1="17" x2="8" y2="17"/></svg>
      Editais
    </button>
    <button class="nav-item" id="nav-upload" onclick="showPage('upload')">
      <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.75" stroke-linecap="round" stroke-linejoin="round"><polyline points="16 16 12 12 8 16"/><line x1="12" y1="12" x2="12" y2="21"/><path d="M20.39 18.39A5 5 0 0 0 18 9h-1.26A8 8 0 1 0 3 16.3"/></svg>
      Novo edital
    </button>
    <button class="nav-item" id="nav-historico" onclick="showPage('historico')">
      <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.75" stroke-linecap="round" stroke-linejoin="round"><circle cx="12" cy="12" r="10"/><polyline points="12 6 12 12 16 14"/></svg>
      HistÃ³rico
    </button>
    <a class="nav-item" href="/status">
      <svg viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.75" stroke-linecap="round" stroke-linejoin="round"><line x1="18" y1="20" x2="18" y2="10"/><line x1="12" y1="20" x2="12" y2="4"/><line x1="6" y1="20" x2="6" y2="14"/></svg>
      Status
    </a>
  </nav>
  <div class="sidebar-footer">
    <div class="quota-label">AnÃ¡lises hoje</div>
    <div class="quota-val" id="quota-val">â€” / 20</div>
    <div class="quota-bar"><div class="quota-fill" id="quota-fill" style="width:0%"></div></div>
    <div class="version-line">{APP_VERSION_LABEL}</div>
    <div class="version-sub">commit {APP_COMMIT_LABEL}</div>
  </div>
</aside>

<div class="main-content" id="main-content"></div>
</div>

<div id="toast"></div>
<input type="file" id="file-input" multiple accept=".pdf,.docx,.xlsx,.xls,.txt" style="display:none">

<script>
var _historico=[];
var _filter='todos';
var _selectedFiles=[];
var _processing=false;

function scoreColor(s){return s>=75?'#166534':s>=50?'#92400E':'#991B1B'}
function scoreLabel(s){return s>=75?'Alta':s>=50?'MÃ©dia':'Baixa'}

function scoreRing(score,size){
  if(!score)return '';
  var r=(size-8)/2,c=2*Math.PI*r,off=c-(score/100)*c,col=scoreColor(score),fs=size<80?13:20;
  return `<svg width="${size}" height="${size}" style="flex-shrink:0"><circle cx="${size/2}" cy="${size/2}" r="${r}" fill="none" stroke="#E5E7EB" stroke-width="4"/><circle cx="${size/2}" cy="${size/2}" r="${r}" fill="none" stroke="${col}" stroke-width="4" stroke-dasharray="${c.toFixed(2)}" stroke-dashoffset="${off.toFixed(2)}" stroke-linecap="round" transform="rotate(-90 ${size/2} ${size/2})"/><text x="${size/2}" y="${size/2}" text-anchor="middle" dominant-baseline="central" font-family="Inter,sans-serif" font-size="${fs}" font-weight="700" fill="${col}">${score}</text></svg>`;
}

function badge(t,cls){return `<span class="badge ${cls}">${t}</span>`}
function viabBadge(s){
  if(!s)return '';
  var cls=s>=75?'badge-success':s>=50?'badge-warn':'badge-danger';
  return badge('Viabilidade '+scoreLabel(s),cls);
}
function fmtDate(ts){
  if(!ts)return '';
  var d=new Date(ts);
  return d.toLocaleDateString('pt-BR',{day:'2-digit',month:'2-digit',hour:'2-digit',minute:'2-digit'});
}
function toast(msg,dur){
  dur=dur||3000;
  var el=document.getElementById('toast'),item=document.createElement('div');
  item.className='toast-item';item.textContent=msg;
  el.appendChild(item);setTimeout(function(){item.remove()},dur);
}
function escHtml(s){return (s||'').replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;')}

function showPage(page,data){
  _filter=page==='editais'?_filter:'todos';
  document.querySelectorAll('.nav-item').forEach(function(el){el.classList.remove('active')});
  var nav=document.getElementById('nav-'+page);
  if(nav)nav.classList.add('active');
  var mc=document.getElementById('main-content');
  if(page==='editais')renderEditaisPage(mc);
  else if(page==='upload')renderUploadPage(mc);
  else if(page==='detalhe')renderDetalhePage(mc,data);
  else if(page==='historico')renderHistoricoPage(mc);
}

function renderEditaisPage(mc){
  var hoje=_historico.filter(function(r){return (r.timestamp||'').split('T')[0]===new Date().toISOString().split('T')[0]});
  var scores=_historico.filter(function(r){return r.score}).map(function(r){return r.score});
  var avgScore=scores.length?Math.round(scores.reduce(function(a,b){return a+b},0)/scores.length):0;
  var filtered=_historico.filter(function(r){
    if(_filter==='alta')return r.score>=75;
    if(_filter==='media')return r.score>=50&&r.score<75;
    if(_filter==='baixa')return r.score>0&&r.score<50;
    return true;
  });
  var filterBtns=['todos','alta','media','baixa'].map(function(f){
    var labels={todos:'Todos',alta:'Viabilidade alta',media:'MÃ©dia',baixa:'Baixa'};
    return `<button class="filter-pill${_filter===f?' active':''}" onclick="setFilter('${f}')">${labels[f]}</button>`;
  }).join('');
  var emptyBtn=_historico.length===0?`<button class="btn btn-primary" style="margin-top:16px" onclick="showPage('upload')">Analisar primeiro edital</button>`:'';
  var cards=filtered.length===0?
    `<div class="empty-state"><svg width="40" height="40" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.5" style="margin:0 auto;display:block;color:var(--fg-4)"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/></svg><div class="empty-title">${_historico.length===0?'Nenhum edital analisado ainda':'Nenhum edital neste filtro'}</div><div class="empty-sub">${_historico.length===0?'FaÃ§a upload de um PDF para comeÃ§ar.':'Tente outro filtro de viabilidade.'}</div>${emptyBtn}</div>`:
    filtered.map(editalCardHTML).join('');
  mc.innerHTML=`<div class="page"><div class="page-header"><div><div class="page-eyebrow">Workspace</div><h1 class="page-title">Editais analisados</h1><p class="page-sub">${_historico.length} edital(is) no histÃ³rico Â· ordenados por data</p></div><button class="btn btn-primary" onclick="showPage('upload')"><svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><polyline points="16 16 12 12 8 16"/><line x1="12" y1="12" x2="12" y2="21"/><path d="M20.39 18.39A5 5 0 0 0 18 9h-1.26A8 8 0 1 0 3 16.3"/></svg>Novo edital</button></div><div class="stats-grid g4"><div class="stat-card"><div class="lbl">AnÃ¡lises hoje</div><div class="val">${hoje.length} <span style="font-size:16px;font-weight:500;color:var(--fg-3)">/ 20</span></div><div class="sub">Reseta Ã  meia-noite</div></div><div class="stat-card"><div class="lbl">Total salvo</div><div class="val">${_historico.length}</div><div class="sub">anÃ¡lises no histÃ³rico</div></div><div class="stat-card"><div class="lbl">Score mÃ©dio</div><div class="val" style="color:${avgScore?scoreColor(avgScore):'var(--fg-4)'}">${avgScore||'â€”'}</div><div class="sub">${avgScore?scoreLabel(avgScore)+' viabilidade':'nenhum calculado'}</div></div><div class="stat-card"><div class="lbl">Alta viabilidade</div><div class="val" style="color:var(--success-700)">${_historico.filter(function(r){return r.score>=75}).length}</div><div class="sub">de ${_historico.length} editais</div></div></div><div class="filter-row">${filterBtns}<div class="filter-right"><button class="btn btn-secondary btn-sm" onclick="showPage('historico')">HistÃ³rico de uso</button></div></div><div class="edital-list">${cards}</div></div>`;
}

function setFilter(f){_filter=f;showPage('editais')}

function editalCardHTML(r){
  var ring=r.score?scoreRing(r.score,64):`<div style="width:64px;height:64px;border-radius:9999px;background:var(--ink-100);display:flex;align-items:center;justify-content:center;font-size:11px;color:var(--fg-4);font-weight:600">â€”</div>`;
  return `<div class="edital-card" onclick="openEdital('${r.id}')">` + ring +
    `<div><div class="edital-meta"><span class="edital-numero">${fmtDate(r.timestamp)}</span>` +
    (r.segmento?badge(r.segmento,'badge-brand'):'') +
    (r.score>=85?badge('PRIORITÃRIO','badge-solid'):'') +
    `</div><div class="edital-objeto">${escHtml(r.objeto||'Sem descriÃ§Ã£o')}</div>` +
    `<div class="edital-orgao">${escHtml(r.orgao||'Ã“rgÃ£o nÃ£o identificado')}</div></div>` +
    `<div class="edital-right"><div class="edital-valor">${escHtml(r.valor||'â€”')}</div>` +
    `<div style="margin-top:6px">${viabBadge(r.score)}</div></div></div>`;
}

async function openEdital(id){
  var r=_historico.find(function(x){return x.id===id});
  if(!r)return;
  if(!r.ficha){
    try{
      var res=await fetch('/historico/'+id);
      var data=await res.json();
      r.ficha=data.ficha;
      if(data.score!=null)r.score=data.score;
    }catch(e){console.error('Erro ao buscar ficha:',e);return;}
  }
  showPage('detalhe',r);
}

function renderUploadPage(mc){
  mc.innerHTML=`<div class="page"><div class="page-header"><div><h1 class="page-title">Novo edital</h1><p class="page-sub">Envie o PDF do edital. A IA extrai objeto, exigÃªncias, valores, prazos e calcula o score de viabilidade.</p></div></div><div class="dropzone" id="dropzone"><div class="dz-icon"><svg width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.75"><polyline points="16 16 12 12 8 16"/><line x1="12" y1="12" x2="12" y2="21"/><path d="M20.39 18.39A5 5 0 0 0 18 9h-1.26A8 8 0 1 0 3 16.3"/></svg></div><div class="dz-title">Arraste o edital ou clique para enviar</div><div class="dz-sub">PDF Â· DOCX Â· XLSX Â· XLS Â· TXT Â· mÃºltiplos arquivos simultÃ¢neos</div></div><div class="file-list" id="file-list"></div><button class="btn btn-primary" id="btn-analisar" style="width:100%;margin-top:20px;justify-content:center;display:none"><svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M12 2L2 7l10 5 10-5-10-5z"/><path d="M2 17l10 5 10-5"/><path d="M2 12l10 5 10-5"/></svg>Analisar edital</button><div style="margin-top:24px;display:grid;grid-template-columns:repeat(3,1fr);gap:12px"><div style="padding:16px;background:var(--bg-subtle);border-radius:12px;border:1px solid var(--border-subtle)"><div style="font-size:13px;font-weight:600;margin-bottom:4px">PDF com texto</div><div style="font-size:12px;color:var(--fg-3)">Use o arquivo original do portal, nÃ£o escaneado.</div></div><div style="padding:16px;background:var(--bg-subtle);border-radius:12px;border:1px solid var(--border-subtle)"><div style="font-size:13px;font-weight:600;margin-bottom:4px">MÃºltiplos arquivos</div><div style="font-size:12px;color:var(--fg-3)">Envie edital + anexos juntos para anÃ¡lise completa.</div></div><div style="padding:16px;background:var(--bg-subtle);border-radius:12px;border:1px solid var(--border-subtle)"><div style="font-size:13px;font-weight:600;margin-bottom:4px">Score automÃ¡tico</div><div style="font-size:12px;color:var(--fg-3)">A IA calcula viabilidade 0â€“100 e lista exigÃªncias.</div></div></div></div>`;
  var dz=document.getElementById('dropzone');
  dz.onclick=function(){document.getElementById('file-input').click()};
  dz.ondragover=function(e){e.preventDefault();dz.classList.add('over')};
  dz.ondragleave=function(){dz.classList.remove('over')};
  dz.ondrop=function(e){e.preventDefault();dz.classList.remove('over');addFiles(e.dataTransfer.files)};
  var btn=document.getElementById('btn-analisar');
  btn.onclick=analisarArquivos;
  renderFileList();
  document.getElementById('file-input').onchange=function(){addFiles(this.files);this.value=''};
}

function handleDrop(e){
  e.preventDefault();
  document.getElementById('dropzone').classList.remove('over');
  addFiles(e.dataTransfer.files);
}
function addFiles(files){for(var i=0;i<files.length;i++)_selectedFiles.push(files[i]);renderFileList()}
function removeFile(i){_selectedFiles.splice(i,1);renderFileList()}
function renderFileList(){
  var fl=document.getElementById('file-list'),btn=document.getElementById('btn-analisar');
  if(!fl)return;
  fl.innerHTML=_selectedFiles.map(function(f,i){
    return `<div class="file-item"><span class="file-name">${escHtml(f.name)}</span><span class="file-size">${Math.round(f.size/1024)} KB</span><button class="file-rm" data-rm="${i}"><svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><line x1="18" y1="6" x2="6" y2="18"/><line x1="6" y1="6" x2="18" y2="18"/></svg></button></div>`;
  }).join('');
  if(btn)btn.style.display=_selectedFiles.length>0?'flex':'none';
  fl.querySelectorAll('[data-rm]').forEach(function(b){b.onclick=function(){removeFile(+b.dataset.rm)}});
}

async function extractPdfText(file){
  if(typeof pdfjsLib==='undefined')return null;
  pdfjsLib.GlobalWorkerOptions.workerSrc='https://cdn.jsdelivr.net/npm/pdfjs-dist@3.11.174/build/pdf.worker.min.js';
  var NL=String.fromCharCode(10);
  var buf=await file.arrayBuffer();
  var pdf=await pdfjsLib.getDocument({data:buf}).promise;
  var maxPgs=Math.min(pdf.numPages,20);
  var parts=[];
  for(var p=1;p<=maxPgs;p++){
    var pg=await pdf.getPage(p);
    var content=await pg.getTextContent();
    var txt=content.items.map(function(it){return it.str}).join(' ');
    if(txt.trim())parts.push(txt);
  }
  return parts.join(NL);
}

async function analisarArquivos(){
  if(_selectedFiles.length===0||_processing)return;
  _processing=true;
  var mc=document.getElementById('main-content');
  mc.innerHTML=`<div class="page"><div class="processing-card"><div class="processing-icon"><svg width="24" height="24" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="1.75"><path d="M12 2L2 7l10 5 10-5-10-5z"/><path d="M2 17l10 5 10-5"/><path d="M2 12l10 5 10-5"/></svg></div><div style="font-size:22px;font-weight:700;letter-spacing:-.015em;margin-bottom:8px">Analisando editalâ€¦</div><div style="font-size:14px;color:var(--fg-2)">Extraindo exigÃªncias, calculando score de viabilidade. AtÃ© 2 minutos.</div><div class="processing-bar"><div class="processing-fill"></div></div></div></div>`;
  try{
    var allPdfs=_selectedFiles.every(function(f){return f.name.toLowerCase().endsWith('.pdf')});
    var resp;
    if(allPdfs&&typeof pdfjsLib!=='undefined'){
      var partes=[];
      for(var i=0;i<_selectedFiles.length;i++){
        var txt=await extractPdfText(_selectedFiles[i]);
        if(!txt||txt.trim().length<100){allPdfs=false;break;}
        var NL2=String.fromCharCode(10);
        partes.push('=== '+_selectedFiles[i].name+' ==='+NL2+txt);
      }
      if(allPdfs){
        var sep=String.fromCharCode(10)+String.fromCharCode(10);
        var textoCompleto=partes.join(sep);
        var res=await fetch('/analisar',{method:'POST',headers:{'Content-Type':'application/json'},body:JSON.stringify({texto:textoCompleto,num_docs:_selectedFiles.length})});
        if(!res.ok){var err=await res.json().catch(function(){return{detail:'Erro desconhecido'}});throw new Error(err.detail||'Erro ao analisar')}
        resp=await res.json();
      }
    }
    if(!resp){
      var fd=new FormData();
      for(var i=0;i<_selectedFiles.length;i++)fd.append('arquivos',_selectedFiles[i]);
      var res=await fetch('/analisar/arquivo',{method:'POST',body:fd});
      if(!res.ok){var err=await res.json().catch(function(){return{detail:'Erro desconhecido'}});throw new Error(err.detail||'Erro ao analisar')}
      resp=await res.json();
    }
    _selectedFiles=[];
    await loadHistorico();
    _processing=false;
    var newest=_historico[0];
    if(newest){newest.ficha=resp.ficha;showPage('detalhe',newest);}else showPage('editais');
    toast('AnÃ¡lise concluÃ­da com sucesso!');
  }catch(e){
    _processing=false;
    toast('Erro: '+e.message,5000);
    showPage('upload');
  }
}

function parseExigencias(ficha){
  var exigs=[],lines=(ficha||'').split('\\n');
  for(var i=0;i<lines.length;i++){
    var m=lines[i].match(/^\\[(ok|warn|fail)\\]\\s*(.+?)(?:\\s*[â€”â€“-]+\\s*(.*))?$/i);
    if(m)exigs.push({status:m[1].toLowerCase(),title:m[2].trim(),detail:(m[3]||'').trim()});
  }
  return exigs;
}

function fichaClean(ficha){
  return (ficha||'')
    .replace(/## Score de Viabilidade[\\s\\S]*?(?=\\n## |\\n*$)/i,'')
    .replace(/## AnÃ¡lise de ExigÃªncias[\\s\\S]*?(?=\\n## |\\n*$)/i,'')
    .trim();
}

var exigIcons={
  ok:'<svg class="exig-icon" width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="#16a34a" stroke-width="2"><path d="M22 11.08V12a10 10 0 1 1-5.93-9.14"/><polyline points="22 4 12 14.01 9 11.01"/></svg>',
  warn:'<svg class="exig-icon" width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="#d97706" stroke-width="2"><path d="M10.29 3.86L1.82 18a2 2 0 0 0 1.71 3h16.94a2 2 0 0 0 1.71-3L13.71 3.86a2 2 0 0 0-3.42 0z"/><line x1="12" y1="9" x2="12" y2="13"/><line x1="12" y1="17" x2="12.01" y2="17"/></svg>',
  fail:'<svg class="exig-icon" width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="#dc2626" stroke-width="2"><circle cx="12" cy="12" r="10"/><line x1="15" y1="9" x2="9" y2="15"/><line x1="9" y1="9" x2="15" y2="15"/></svg>'
};

function renderDetalhePage(mc,r){
  var exigs=parseExigencias(r.ficha);
  var exigHTML=exigs.length>0?
    `<div style="margin-top:40px"><h3 style="font-size:18px;font-weight:700;letter-spacing:-.015em;margin-bottom:6px">AnÃ¡lise de exigÃªncias</h3><p style="font-size:14px;color:var(--fg-2);margin-bottom:16px">${exigs.length} exigÃªncia(s) Â· ${exigs.filter(function(e){return e.status==='ok'}).length} ok Â· ${exigs.filter(function(e){return e.status==='warn'}).length} atenÃ§Ã£o Â· ${exigs.filter(function(e){return e.status==='fail'}).length} restritiva(s)</p>` +
    exigs.map(function(e){
      return `<div class="exig-item exig-${e.status}">${exigIcons[e.status]}<div><div class="exig-title">${escHtml(e.title)}</div>${e.detail?`<div class="exig-detail">${escHtml(e.detail)}</div>`:''}</div></div>`;
    }).join('') + `</div>` : '';

  var fichaHtml=typeof marked!=='undefined'?marked.parse(fichaClean(r.ficha)):
    `<pre style="white-space:pre-wrap;font-size:13px">${escHtml(fichaClean(r.ficha))}</pre>`;
  var arquivosHtml=(r.arquivos||[]).slice().sort(function(a,b){return (a.ordem||0)-(b.ordem||0)}).length>0?
    `<div style="margin-top:40px"><h3 style="font-size:18px;font-weight:700;letter-spacing:-.015em;margin-bottom:8px">Arquivos originais</h3><p style="font-size:14px;color:var(--fg-2);margin-bottom:16px">${(r.arquivos||[]).length} arquivo(s) preservado(s) junto desta análise.</p><div style="display:flex;flex-direction:column;gap:10px">${(r.arquivos||[]).slice().sort(function(a,b){return (a.ordem||0)-(b.ordem||0)}).map(function(a){var url='/historico/'+r.id+'/arquivos/'+a.id;return `<div class="stat-card" style="padding:14px 16px;display:flex;align-items:center;justify-content:space-between;gap:12px"><div style="min-width:0"><div style="font-size:13px;font-weight:600;white-space:nowrap;overflow:hidden;text-overflow:ellipsis">${escHtml(a.arquivo||'arquivo')}</div><div style="font-size:12px;color:var(--fg-3);margin-top:4px">${escHtml(a.mime_type||'application/octet-stream')} · ${a.tamanho_bytes?Math.round(a.tamanho_bytes/1024)+' KB':'—'}</div></div><a class="btn btn-secondary btn-sm" href="${url}">Baixar</a></div>`;}).join('')}</div></div>` : '';

  mc.innerHTML=
    `<div class="page"><button class="back-btn" onclick="showPage('editais')"><svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="m15 18-6-6 6-6"/></svg>Voltar para editais</button>` +
    `<div style="display:flex;align-items:flex-start;justify-content:space-between;gap:32px;margin-bottom:32px"><div style="flex:1"><div style="display:flex;gap:8px;align-items:center;margin-bottom:12px;flex-wrap:wrap">` +
    (r.timestamp?`<span style="font-family:var(--font-mono);font-size:12px;color:var(--fg-3)">${fmtDate(r.timestamp)}</span>`:'') +
    (r.segmento?badge(r.segmento,'badge-brand'):'') +
    (r.score>=85?badge('PRIORITÃRIO','badge-solid'):'') +
    `</div><h1 style="font-size:28px;font-weight:700;letter-spacing:-.025em;line-height:1.2;margin-bottom:8px">${escHtml(r.objeto||'AnÃ¡lise de edital')}</h1>` +
    `<p style="font-size:15px;color:var(--fg-2)">${escHtml(r.orgao||'')}</p></div>` +
    (r.score?`<div style="flex-shrink:0;text-align:center">${scoreRing(r.score,120)}<div style="margin-top:8px;font-size:13px;font-weight:600;color:${scoreColor(r.score)}">Viabilidade ${scoreLabel(r.score)}</div></div>`:'') +
    `</div>` +
    `<div class="stats-grid g4" style="margin-bottom:40px"><div class="stat-card"><div class="lbl">Valor estimado</div><div class="val" style="font-size:18px">${escHtml(r.valor||'â€”')}</div></div><div class="stat-card"><div class="lbl">Segmento</div><div class="val" style="font-size:18px">${escHtml(r.segmento||'â€”')}</div></div><div class="stat-card"><div class="lbl">Score</div><div class="val" style="color:${r.score?scoreColor(r.score):'var(--fg-4)'}">${r.score||'â€”'}</div><div class="sub">${r.score?scoreLabel(r.score)+' viabilidade':''}</div></div><div class="stat-card"><div class="lbl">Analisado em</div><div class="val" style="font-size:16px">${r.timestamp?new Date(r.timestamp).toLocaleDateString('pt-BR'):'â€”'}</div></div></div>` +
    exigHTML +
    arquivosHtml +
    `<div style="margin-top:40px"><h3 style="font-size:18px;font-weight:700;letter-spacing:-.015em;margin-bottom:16px">Ficha completa</h3><div class="ficha-content">${fichaHtml}</div></div>` +
    `<div class="action-row"><button class="btn btn-secondary" id="btn-print"><svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><polyline points="6 9 6 2 18 2 18 9"/><path d="M6 18H4a2 2 0 0 1-2-2v-5a2 2 0 0 1 2-2h16a2 2 0 0 1 2 2v5a2 2 0 0 1-2 2h-2"/><rect x="6" y="14" width="12" height="8"/></svg>Imprimir / PDF</button><button class="btn btn-secondary" id="btn-copy"><svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><rect x="9" y="9" width="13" height="13" rx="2"/><path d="M5 15H4a2 2 0 0 1-2-2V4a2 2 0 0 1 2-2h9a2 2 0 0 1 2 2v1"/></svg>Copiar ficha</button><button class="btn btn-primary"><svg width="16" height="16" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M12 2L2 7l10 5 10-5-10-5z"/><path d="M2 17l10 5 10-5"/><path d="M2 12l10 5 10-5"/></svg>Gerar proposta</button></div></div>`;
  document.getElementById('btn-print').onclick=function(){window.print()};
  document.getElementById('btn-copy').onclick=function(){copiarFicha(r.id)};
  mc.querySelector('.btn-primary:last-child').onclick=function(){toast('Em breve: geraÃ§Ã£o de proposta por IA.')};
}

async function copiarFicha(id){
  var r=_historico.find(function(x){return x.id===id});
  if(!r)return;
  if(!r.ficha){
    try{var res=await fetch('/historico/'+id);var data=await res.json();r.ficha=data.ficha;}
    catch(e){toast('Erro ao copiar');return;}
  }
  navigator.clipboard.writeText(r.ficha||'').then(function(){toast('Ficha copiada!')}).catch(function(){toast('Erro ao copiar')});
}

function renderHistoricoPage(mc){
  var byDay={};
  _historico.forEach(function(r){
    var d=(r.timestamp||'').split('T')[0];
    if(!d)return;
    if(!byDay[d])byDay[d]=[];
    byDay[d].push(r);
  });
  var days=Object.keys(byDay).sort().reverse();
  var scores=_historico.filter(function(r){return r.score}).map(function(r){return r.score});
  var avg=scores.length?Math.round(scores.reduce(function(a,b){return a+b},0)/scores.length):0;
  var tableRows=days.map(function(d){
    var items=byDay[d];
    var s=items.filter(function(r){return r.score}).map(function(r){return r.score});
    var davg=s.length?Math.round(s.reduce(function(a,b){return a+b},0)/s.length):0;
    var altas=items.filter(function(r){return r.score>=75}).length;
    return `<tr><td style="font-family:var(--font-mono)">${new Date(d+'T12:00:00').toLocaleDateString('pt-BR')}</td><td>${items.length}</td><td style="font-weight:600;color:${davg?scoreColor(davg):'var(--fg-4)'}">${davg||'â€”'}</td><td style="color:var(--success-700);font-weight:600">${altas}</td></tr>`;
  }).join('');
  var seg={};
  _historico.forEach(function(r){var s=r.segmento||'Outros';seg[s]=(seg[s]||0)+1});
  var segRows=Object.entries(seg).sort(function(a,b){return b[1]-a[1]}).map(function(e){
    return `<div style="display:flex;justify-content:space-between;align-items:center;padding:10px 14px;background:#fff;border:1px solid var(--border);border-radius:8px;font-size:13px"><span style="color:var(--fg-2)">${escHtml(e[0])}</span><span style="font-weight:600;color:var(--brand-600);background:var(--brand-50);padding:2px 10px;border-radius:9999px;font-size:12px">${e[1]}</span></div>`;
  }).join('');

  mc.innerHTML=`<div class="page"><div class="page-header"><div><h1 class="page-title">HistÃ³rico de uso</h1><p class="page-sub">Todas as anÃ¡lises agrupadas por dia</p></div><button class="btn btn-secondary" onclick="showPage('editais')">â† Editais</button></div><div class="stats-grid g3"><div class="stat-card"><div class="lbl">Total de anÃ¡lises</div><div class="val">${_historico.length}</div><div class="sub">no histÃ³rico</div></div><div class="stat-card"><div class="lbl">Score mÃ©dio</div><div class="val" style="color:${avg?scoreColor(avg):'var(--fg-4)'}">${avg||'â€”'}</div><div class="sub">${avg?scoreLabel(avg)+' viabilidade':''}</div></div><div class="stat-card"><div class="lbl">Dias com anÃ¡lises</div><div class="val">${days.length}</div><div class="sub">dias distintos</div></div></div>${days.length===0?'<div class="empty-state"><div class="empty-title">Nenhuma anÃ¡lise</div></div>':`<div class="hist-table"><table><thead><tr><th>Dia</th><th>AnÃ¡lises</th><th>Score mÃ©dio</th><th>Alta viabilidade</th></tr></thead><tbody>${tableRows}</tbody></table></div>`}<div style="margin-top:32px"><h3 style="font-size:16px;font-weight:700;margin-bottom:16px">Por segmento</h3><div style="display:flex;flex-direction:column;gap:8px">${segRows||'<p style="font-size:14px;color:var(--fg-3)">Nenhum dado ainda.</p>'}</div></div></div>`;
}

async function loadHistorico(){
  try{
    var res=await fetch('/historico'),data=await res.json();
    _historico=data.historico||[];
    var hoje=_historico.filter(function(r){return (r.timestamp||'').split('T')[0]===new Date().toISOString().split('T')[0]}).length;
    var qv=document.getElementById('quota-val'),qf=document.getElementById('quota-fill');
    if(qv)qv.textContent=hoje+' / 20';
    if(qf)qf.style.width=Math.min(100,(hoje/20)*100)+'%';
  }catch(e){console.error('Erro ao carregar histÃ³rico:',e)}
}

async function initApp(){await loadHistorico();showPage('editais')}
initApp();
</script>
</body>
</html>"""


class AnalisarRequest(BaseModel):
    texto: str
    num_docs: int = 2
    modo: str = "auto"  # "auto" | "parser" | "ia"


class AnalisarResponse(BaseModel):
    ficha: str


def _obter_ocr_engine():
    global _OCR_ENGINE
    if not OCR_HABILITADO or RapidOCR is None:
        return None
    if _OCR_ENGINE is None:
        try:
            _OCR_ENGINE = RapidOCR()
        except Exception as e:
            logger.exception("Erro ao inicializar motor OCR", extra={"request_id": "-"})
            _OCR_ENGINE = False
            return None
    return None if _OCR_ENGINE is False else _OCR_ENGINE


def _texto_precisa_ocr(texto: str) -> bool:
    texto_limpo = re.sub(r"\s+", "", texto or "")
    return OCR_HABILITADO and len(texto_limpo) < OCR_MIN_CHAR


def _ocr_resultado_para_texto(resultado) -> str:
    if not resultado:
        return ""

    if isinstance(resultado, tuple) and resultado:
        resultado = resultado[0]

    linhas = []
    for item in resultado:
        if not item:
            continue
        try:
            box, texto, *_ = item
        except Exception:
            continue
        texto = str(texto or "").strip()
        if not texto:
            continue
        x = 0
        y = 0
        try:
            if box:
                xs = [p[0] for p in box]
                ys = [p[1] for p in box]
                x = min(xs)
                y = min(ys)
        except Exception:
            pass
        linhas.append((y, x, texto))

    linhas.sort(key=lambda item: (item[0], item[1]))
    return "\n".join(texto for _, _, texto in linhas).strip()


def _ocr_pagina_pdf(doc, pagina_idx: int) -> str:
    if fitz is None or np is None:
        return ""
    engine = _obter_ocr_engine()
    if engine is None:
        return ""

    try:
        pagina = doc[pagina_idx]
        matriz = fitz.Matrix(OCR_DPI / 72.0, OCR_DPI / 72.0)
        pix = pagina.get_pixmap(matrix=matriz, alpha=False)
        imagem = np.frombuffer(pix.samples, dtype=np.uint8)
        if pix.n == 1:
            imagem = imagem.reshape(pix.height, pix.width)
            imagem = np.stack([imagem] * 3, axis=-1)
        else:
            imagem = imagem.reshape(pix.height, pix.width, pix.n)
            if imagem.shape[2] > 3:
                imagem = imagem[:, :, :3]
            elif imagem.shape[2] == 1:
                imagem = np.repeat(imagem, 3, axis=2)
        resultado = engine(imagem)
        return _ocr_resultado_para_texto(resultado)
    except Exception as e:
        logger.warning(
            "Erro na pagina OCR %s: %s",
            pagina_idx + 1,
            e,
            extra={"request_id": "-"},
        )
        return ""


def _extrair_texto_pdf(conteudo: bytes) -> str:
    pdf_ocr = None
    ocr_usado = False
    arquivo_grande = len(conteudo) > OCR_MAX_FILE_BYTES
    if arquivo_grande:
        logger.info(
            "Arquivo %s KB > limite %s KB - OCR desativado.",
            len(conteudo) // 1024,
            OCR_MAX_FILE_BYTES // 1024,
            extra={"request_id": "-"},
        )
    try:
        if fitz is not None and OCR_HABILITADO and not arquivo_grande:
            pdf_ocr = fitz.open(stream=conteudo, filetype="pdf")
        with pdfplumber.open(io.BytesIO(conteudo)) as pdf:
            partes = []
            limite_paginas = MAX_PAGINAS_GRANDES if arquivo_grande else len(pdf.pages)
            for indice, p in enumerate(pdf.pages[:limite_paginas]):
                texto = p.extract_text() or ""
                if not texto.strip():
                    try:
                        words = p.extract_words(x_tolerance=5, y_tolerance=5)
                        if words:
                            texto = " ".join(w["text"] for w in words)
                    except Exception:
                        pass

                if _texto_precisa_ocr(texto) and pdf_ocr is not None and indice < OCR_MAX_PAGINAS:
                    texto_ocr = _ocr_pagina_pdf(pdf_ocr, indice)
                    if texto_ocr.strip():
                        texto = texto_ocr
                        ocr_usado = True

                partes.append(texto)

        resultado = "\n".join(partes).strip()
        if not resultado:
            # Ãºltimo recurso: OCR em sequÃªncia quando a extraÃ§Ã£o nativa vier vazia.
            if pdf_ocr is not None:
                partes_ocr = []
                for indice in range(min(len(pdf_ocr), OCR_MAX_PAGINAS)):
                    texto_ocr = _ocr_pagina_pdf(pdf_ocr, indice)
                    if texto_ocr.strip():
                        partes_ocr.append(texto_ocr)
                resultado = "\n".join(partes_ocr).strip()
                if resultado:
                    ocr_usado = True

        if not resultado:
            # Ãºltimo recurso: pdfminer.six diretamente
            try:
                from pdfminer.high_level import extract_text as _pm_et
                resultado = (_pm_et(io.BytesIO(conteudo)) or "").strip()
            except Exception:
                pass

        if ocr_usado:
            logger.info("OCR aplicado durante o upload do PDF.", extra={"request_id": "-"})
        return resultado
    finally:
        if pdf_ocr is not None:
            try:
                pdf_ocr.close()
            except Exception:
                pass


def _texto_de_odt(conteudo: bytes) -> str:
    try:
        with zipfile.ZipFile(io.BytesIO(conteudo)) as zf:
            with zf.open("content.xml") as fh:
                root = ET.fromstring(fh.read())
    except Exception as e:
        print(f"[ODT] Erro ao ler arquivo: {e}")
        return ""

    ns = {
        "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
    }
    partes = []
    for tag in (".//text:h", ".//text:p"):
        for node in root.findall(tag, ns):
            trecho = "".join(node.itertext()).strip()
            if trecho:
                partes.append(trecho)
    return "\n".join(partes).strip()


def _texto_de_docx(conteudo: bytes) -> str:
    try:
        with zipfile.ZipFile(io.BytesIO(conteudo)) as zf:
            with zf.open("word/document.xml") as fh:
                root = ET.fromstring(fh.read())
    except Exception as e:
        print(f"[DOCX] Erro ao ler arquivo: {e}")
        return ""

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    partes = []
    for node in root.findall(".//w:t", ns):
        trecho = (node.text or "").strip()
        if trecho:
            partes.append(trecho)
    return "\n".join(partes).strip()


def extrair_texto(nome: str, conteudo: bytes) -> str:
    nome_lower = nome.lower()
    if nome_lower.endswith(".pdf"):
        return _extrair_texto_pdf(conteudo)
    if nome_lower.endswith(".docx"):
        texto_xml = _texto_de_docx(conteudo)
        if texto_xml.strip():
            return texto_xml
        doc = docx.Document(io.BytesIO(conteudo))
        return "\n".join(p.text for p in doc.paragraphs)
    if nome_lower.endswith((".xlsx", ".xls")):
        wb = openpyxl.load_workbook(io.BytesIO(conteudo), data_only=True)
        linhas = []
        for nome_aba in wb.sheetnames:
            ws = wb[nome_aba]
            linhas.append(f"[Aba: {nome_aba}]")
            for row in ws.iter_rows(values_only=True):
                linha = "\t".join("" if c is None else str(c) for c in row)
                if linha.strip():
                    linhas.append(linha)
        return "\n".join(linhas)
    if nome_lower.endswith(".odt"):
        return _texto_de_odt(conteudo)
    return conteudo.decode("utf-8", errors="replace")


def extrair_texto_com_tipo(nome: str, conteudo: bytes) -> tuple[str, str]:
    nome_lower = nome.lower()
    if nome_lower.endswith(".pdf"):
        return "pdf", extrair_texto(nome, conteudo)
    if nome_lower.endswith(".docx"):
        return "docx", extrair_texto(nome, conteudo)
    if nome_lower.endswith((".xlsx", ".xls")):
        return "planilha", extrair_texto(nome, conteudo)
    if nome_lower.endswith(".odt"):
        return "odt", extrair_texto(nome, conteudo)
    if nome_lower.endswith((".txt", ".md", ".csv")):
        return "texto", extrair_texto(nome, conteudo)
    return "outro", extrair_texto(nome, conteudo)


def montar_texto_caso_raw(arquivos: list[tuple[str, bytes]]) -> tuple[str, list[dict[str, str]]]:
    textos = []
    meta = []
    for nome, conteudo in arquivos:
        tipo, texto = extrair_texto_com_tipo(nome, conteudo)
        texto = (texto or "").strip()
        meta.append({"arquivo": nome, "tipo": tipo, "chars": str(len(texto))})
        if texto:
            textos.append((nome, tipo, texto))

    partes = []
    for nome, tipo, txt in textos:
        partes.append(f"=== {nome} [{tipo}] ===\n{txt}")
    return "\n\n".join(partes), meta


def _normalizar_espacos(texto: str) -> str:
    return re.sub(r"\s+", " ", (texto or "")).strip()


def _score_principal(nome: str, tipo: str, texto: str) -> int:
    n = (nome or "").lower()
    t = _normalizar_espacos(texto[:25000]).lower()

    score = 0
    # Tipo do arquivo (preferir texto “narrativo”)
    if tipo == "pdf":
        score += 8
    elif tipo == "docx":
        score += 7
    elif tipo == "odt":
        score += 5
    elif tipo == "texto":
        score += 4
    elif tipo == "planilha":
        score -= 3

    # Nome do arquivo
    if "edital" in n:
        score += 8
    if "preg" in n or "pe" in n:
        score += 3
    if "termo" in n or "tr" in n or "etp" in n:
        score += 2
    if "minuta" in n or "contrato" in n:
        score -= 1
    if "anexo" in n:
        score -= 1
    if "relacao" in n or "itens" in n:
        score -= 1

    # Sinais de cabeçalho de edital
    for kw, w in (
        ("pregão eletrônico", 10),
        ("pregao eletronico", 10),
        ("processo", 5),
        ("uasg", 6),
        ("contratante", 4),
        ("objeto", 3),
        ("valor total da contratação", 6),
        ("data da sessão pública", 6),
        ("critério de julgamento", 5),
        ("criterio de julgamento", 5),
    ):
        if kw in t:
            score += w

    return score


def _extrair_cabecalho_hint(texto_principal: str) -> str:
    """
    Gera um preâmbulo estruturado (curto) para ajudar o parser em casos compostos.
    Não substitui o texto bruto; apenas dá pistas com campos nomeados.
    """
    t = _normalizar_espacos(texto_principal[:60000])
    if not t:
        return ""

    def first(pat: str) -> str:
        m = re.search(pat, t, re.IGNORECASE)
        return m.group(1).strip() if m else ""

    # Casos comuns em editais federais: "PREGÃO ELETRÔNICO Nº 90009/2026"
    numero = first(r"\b((?:preg[aã]o\s+eletr[oô]nico|edital|processo(?:\s+administrativo)?)\s*(?:n[ºo°]\s*)?[\.: -]*\s*\d{1,6}[./-]\d{4})\b")
    orgao = first(r"\b((?:TRIBUNAL REGIONAL DO TRABALHO[^.]{0,140}|CONSELHO REGIONAL DE PSICOLOGIA[^.]{0,140}|PREFEITURA MUNICIPAL[^.]{0,140}|INSTITUTO INTERAMERICANO[^.]{0,160}))\b")
    valor = first(r"\b(R\$\s*[\d\.]{1,3}(?:\.[\d]{3})*(?:,\d{2})?)\b")
    data = first(r"\b(\d{2}/\d{2}/\d{4})\b")
    modalidade = "Pregão Eletrônico" if re.search(r"\bpreg[aã]o\b", t, re.IGNORECASE) and re.search(r"\beletr[oô]nico\b", t, re.IGNORECASE) else ""
    criterio = first(r"\b((?:menor pre[cç]o(?:\s+global|\s+por\s+item)?|maior desconto|t[eé]cnica e pre[cç]o|melhor t[eé]cnica))\b")

    linhas = ["## FICHA DE LICITAÇÃO (HINT)"]
    if numero:
        linhas.append(f"- Número do edital/processo: {numero}")
    if orgao:
        linhas.append(f"- Órgão: {orgao}")
    if modalidade:
        linhas.append(f"- Modalidade: {modalidade}")
    if criterio:
        linhas.append(f"- Critério de julgamento: {criterio}")
    if valor:
        linhas.append(f"- Valor estimado total: {valor}")
    if data:
        linhas.append(f"- Data de abertura: {data}")

    if len(linhas) == 1:
        return ""
    return "\n".join(linhas) + "\n\n"


def montar_texto_caso_classificado_raw(arquivos: list[tuple[str, bytes]]) -> tuple[str, list[dict[str, str]]]:
    """
    Monta o texto do caso priorizando o arquivo principal (classificado) e
    adicionando um cabeçalho (hint) com campos nomeados do principal.
    """
    itens = []
    meta = []
    for nome, conteudo in arquivos:
        tipo, texto = extrair_texto_com_tipo(nome, conteudo)
        texto = (texto or "").strip()
        meta.append({"arquivo": nome, "tipo": tipo, "chars": str(len(texto))})
        itens.append({"arquivo": nome, "tipo": tipo, "texto": texto})

    if not itens:
        return "", meta

    # escolhe principal
    principal = max(itens, key=lambda it: _score_principal(it["arquivo"], it["tipo"], it["texto"]))

    # ordena: principal primeiro, depois os demais por score desc
    outros = [it for it in itens if it is not principal]
    outros.sort(key=lambda it: _score_principal(it["arquivo"], it["tipo"], it["texto"]), reverse=True)
    ordenados = [principal] + outros

    cabecalho = _extrair_cabecalho_hint(principal["texto"])
    partes = []
    for it in ordenados:
        if not it["texto"]:
            continue
        partes.append(f"=== {it['arquivo']} [{it['tipo']}] ===\n{it['texto']}")

    return (cabecalho + "\n\n".join(partes)).strip(), meta


def montar_texto_caso(arquivos: list[UploadFile]) -> tuple[str, list[dict[str, str]]]:
    return montar_texto_caso_raw([(a.filename, a.file.read()) for a in arquivos])


def montar_texto_caso_classificado(arquivos: list[UploadFile]) -> tuple[str, list[dict[str, str]]]:
    return montar_texto_caso_classificado_raw([(a.filename, a.file.read()) for a in arquivos])


# ── Histórico de análises ────────────────────────────────────────────────────
HISTORICO_FILE = Path("historico.json")
_DATABASE_URL = os.getenv("DATABASE_URL")

def _db_conn():
    import psycopg2
    return psycopg2.connect(_DATABASE_URL)


def _normalizar_meta_arquivo(nome: str, conteudo: bytes, ordem: int, extra: Optional[dict] = None) -> dict:
    meta = {
        "id": uuid.uuid4().hex[:12],
        "arquivo": nome,
        "mime_type": mimetypes.guess_type(nome)[0] or "application/octet-stream",
        "tamanho_bytes": len(conteudo),
        "sha256": hashlib.sha256(conteudo).hexdigest(),
        "ordem": ordem,
    }
    if extra:
        for chave, valor in extra.items():
            if valor is not None:
                meta[chave] = valor
    return meta

def _init_db():
    if not _DATABASE_URL:
        return
    try:
        with _db_conn() as conn:
            with conn.cursor() as cur:
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS historico (
                        id TEXT PRIMARY KEY,
                        dados JSONB NOT NULL
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS historico_arquivos (
                        id TEXT PRIMARY KEY,
                        analise_id TEXT NOT NULL REFERENCES historico(id) ON DELETE CASCADE,
                        ordem INTEGER NOT NULL,
                        nome_original TEXT NOT NULL,
                        mime_type TEXT,
                        tamanho_bytes INTEGER,
                        sha256 TEXT,
                        tipo_extraido TEXT,
                        chars_extraidos INTEGER,
                        conteudo BYTEA NOT NULL,
                        criado_em TIMESTAMPTZ DEFAULT now()
                    )
                """)
                cur.execute("""
                    CREATE INDEX IF NOT EXISTS idx_historico_arquivos_analise_id
                    ON historico_arquivos (analise_id, ordem)
                """)
            conn.commit()
    except Exception as e:
        logger.exception("Erro ao inicializar tabela do hist?rico", extra={"request_id": "-"})

def _carregar_historico() -> list:
    if _DATABASE_URL:
        try:
            with _db_conn() as conn:
                with conn.cursor() as cur:
                    cur.execute("SELECT dados FROM historico ORDER BY dados->>'timestamp' DESC")
                    rows = cur.fetchall()
                    historico = [r[0] for r in rows] if rows else []
                    if historico:
                        cur.execute("""
                            SELECT
                                id,
                                analise_id,
                                ordem,
                                nome_original,
                                mime_type,
                                tamanho_bytes,
                                sha256,
                                tipo_extraido,
                                chars_extraidos
                            FROM historico_arquivos
                            ORDER BY analise_id, ordem, criado_em
                        """)
                        anexos: dict[str, list[dict]] = {}
                        for arq in cur.fetchall():
                            analise_id = arq[1]
                            anexos.setdefault(analise_id, []).append({
                                "id": arq[0],
                                "arquivo": arq[3],
                                "mime_type": arq[4],
                                "tamanho_bytes": arq[5],
                                "sha256": arq[6],
                                "ordem": arq[2],
                                "tipo_extraido": arq[7],
                                "chars_extraidos": arq[8],
                            })
                        for item in historico:
                            item["arquivos"] = anexos.get(item.get("id"), [])
                    return historico
        except Exception as e:
            logger.exception("Erro ao carregar hist?rico do banco", extra={"request_id": "-"})
    try:
        if HISTORICO_FILE.exists():
            return json.loads(HISTORICO_FILE.read_text(encoding="utf-8"))
    except Exception:
        pass
    return []

_init_db()
_historico: list = _carregar_historico()

# migra automaticamente do arquivo para o banco se o banco estiver vazio
def _migrar_se_necessario():
    if not _DATABASE_URL or not _historico:
        return
    try:
        with _db_conn() as conn:
            with conn.cursor() as cur:
                cur.execute("SELECT COUNT(*) FROM historico")
                count = cur.fetchone()[0]
        if count == 0:
            _salvar_historico()
            logger.info("Migra??o do hist?rico: %s an?lise(s) importadas do arquivo.", len(_historico), extra={"request_id": "-"})
    except Exception as e:
        logger.exception("Erro na migra??o do hist?rico", extra={"request_id": "-"})

_migrar_se_necessario()

def _salvar_historico():
    if _DATABASE_URL:
        try:
            with _db_conn() as conn:
                with conn.cursor() as cur:
                    ids_atuais = []
                    for item in _historico:
                        ids_atuais.append(item.get("id"))
                        item_db = {k: v for k, v in item.items() if k != "arquivos"}
                        cur.execute("""
                            INSERT INTO historico (id, dados)
                            VALUES (%s, %s::jsonb)
                            ON CONFLICT (id) DO UPDATE SET dados = EXCLUDED.dados
                        """, (item.get("id"), json.dumps(item_db, ensure_ascii=False)))
                    if ids_atuais:
                        cur.execute("DELETE FROM historico WHERE NOT (id = ANY(%s::text[]))", (ids_atuais,))
                    else:
                        cur.execute("DELETE FROM historico_arquivos")
                        cur.execute("DELETE FROM historico")
                conn.commit()
            return
        except Exception as e:
            logger.exception("Erro ao salvar historico", extra={"request_id": "-"})

    try:
        HISTORICO_FILE.write_text(
            json.dumps(_historico, ensure_ascii=False, indent=2), encoding="utf-8"
        )
    except Exception:
        pass

_SEGMENTOS = [
    ("SaÃºde",                  ["saÃºde", "saude", "mÃ©dic", "medic", "hospital", "medicament", "ubs", "enfermagem", "cirÃºrgic", "farmÃ¡c", "farmac", "ambulatorial"]),
    ("EducaÃ§Ã£o",               ["escola", "educaÃ§Ã£o", "educacao", "pedagÃ³g", "pedagogic", "didÃ¡tic", "ensino", "aluno", "professor", "material escolar", "creche"]),
    ("Obras e Infraestrutura", ["obras", "construÃ§Ã£o", "construcao", "reforma", "paviment", "infraestrutura", "engenharia", "elÃ©tric", "eletric", "hidrÃ¡ulic", "hidraulic", "saneamento"]),
    ("AlimentaÃ§Ã£o",            ["aliment", "merenda", "refeiÃ§Ã£o", "refeicao", "gÃªneros alimentÃ­c", "generos aliment", "nutri", "cozinha", "marmita"]),
    ("Tecnologia e TI",        ["software", "hardware", "computador", "informÃ¡tica", "informatica", "sistema", "licenÃ§a", "servidor", " ti ", "tecnologia da informaÃ§Ã£o", "impressora"]),
    ("Transporte",             ["veÃ­culo", "veiculo", "frota", "combustÃ­vel", "combustivel", "Ã´nibus", "onibus", "manutenÃ§Ã£o veicular", "locaÃ§Ã£o de veÃ­culo", "locacao de veiculo"]),
    ("Viagens e Passagens",    ["passagem aÃ©rea", "passagem aerea", "passagem Ã¡rea", "bilhete aÃ©reo", "bilhete aereo", "aÃ©reo", "aereo", "aÃ©rea", "aerea", "aviaÃ§Ã£o", "aviacao", "companhia aÃ©rea", "companhia aerea", "passagem", "hospedagem", "diÃ¡ria", "diaria", "hotel", "viagem"]),
    ("Eventos e CapacitaÃ§Ã£o",  ["evento", "congresso", "capacitaÃ§Ã£o", "capacitacao", "treinamento", "curso", "palestra", "cerimÃ´nia", "cerimonia"]),
    ("Limpeza e ConservaÃ§Ã£o",  ["limpeza", "higien", "conservaÃ§Ã£o predial", "conservacao predial", "jardinagem", "desinfeÃ§Ã£o", "desinfecao", "asseio", "zeladoria"]),
    ("MobiliÃ¡rio e EscritÃ³rio",["mobiliÃ¡rio", "mobiliario", "mobÃ­lia", "mobilia", "escritÃ³rio", "escritorio", "papel", "caneta", "grampe", "cadeira", "mesa", "material de escritÃ³rio", "material de escritorio"]),
    ("SeguranÃ§a",              ["seguranÃ§a", "seguranca", "vigilÃ¢ncia", "vigilancia", "monitoramento", "cÃ¢mera", "camera", "cctv", "alarme", "portaria"]),
]

def detectar_segmento(texto: str) -> str:
    t = texto.lower()
    for nome, palavras in _SEGMENTOS:
        if any(p in t for p in palavras):
            return nome
    return "Outros"

def extrair_campo(ficha: str, campo: str) -> str:
    # padrÃ£o markdown: | **Campo** | valor |
    m = re.search(rf'\|\s*\*\*{re.escape(campo)}\*\*\s*\|\s*([^|\n]+)', ficha)
    if m:
        return m.group(1).strip()
    # padrÃ£o texto puro: "Campo: valor" ou "Campo    valor"
    m = re.search(rf'(?:^|\n)\s*{re.escape(campo)}\s*[:\t|]+\s*([^\n]+)', ficha, re.IGNORECASE)
    if m:
        return m.group(1).strip()[:150]
    return "NÃ£o informado"

def extrair_objeto(ficha: str) -> str:
    # seÃ§Ã£o markdown
    m = re.search(r'## Objeto\s*\n+(.+?)(?:\n##|\Z)', ficha, re.DOTALL)
    if m:
        return m.group(1).strip()[:200]
    # seÃ§Ã£o texto puro
    m = re.search(r'Objeto\s*\n+(.+?)(?:\n\n|\Z)', ficha, re.DOTALL)
    if m:
        return m.group(1).strip()[:200]
    return "NÃ£o informado"

def extrair_score(ficha: str) -> int:
    m = re.search(r'\*\*Score:\*\*\s*(\d+)', ficha)
    if m:
        try:
            return max(0, min(100, int(m.group(1))))
        except Exception:
            pass
    return 0

def _eh_ficha(texto: str) -> bool:
    # normaliza espaÃ§os e coloca em maiÃºsculas
    t = " ".join(texto.upper().split())
    # match direto (com ou sem acento)
    if "FICHA DE LICITAÃ‡" in t or "FICHA DE LICITAC" in t:
        return True
    # "FICHA" + "LICITA" em qualquer lugar (cobre encoding quebrado)
    if "FICHA" in t and "LICITA" in t:
        return True
    # fallback: presenÃ§a de vÃ¡rios campos tÃ­picos de ficha
    campos_tipicos = [
        "VALOR ESTIMADO", "DOCUMENTOS DE HABILITA", "ITENS A COTAR",
        "MODALIDADE", "VIGÃŠNCIA", "VIGENCIA", "CRITÃ‰RIO", "CRITERIO",
        "PRAZO DE PAGAMENTO", "CONTATO DO", "ABERTURA",
    ]
    return sum(1 for c in campos_tipicos if c in t) >= 3

def registrar_analise(ficha: str):
    registro = {
        "id":        uuid.uuid4().hex[:10],
        "timestamp": datetime.datetime.now().isoformat(timespec="seconds"),
        "orgao":     extrair_campo(ficha, "Ã“rgÃ£o"),
        "valor":     extrair_campo(ficha, "Valor Estimado Total"),
        "objeto":    extrair_objeto(ficha),
        "segmento":  detectar_segmento(ficha),
        "score":     extrair_score(ficha),
        "ficha":     ficha,
    }
    _historico.insert(0, registro)
    if len(_historico) > 500:
        _historico.pop()
    _salvar_historico()

def _reclassificar_historico():
    mudou = False
    for r in _historico:
        novo = detectar_segmento(r.get("ficha", r.get("objeto", "")))
        if r.get("segmento") != novo:
            r["segmento"] = novo
            mudou = True
    if mudou:
        _salvar_historico()

_reclassificar_historico()


def _normalizar_meta_extracao(meta_arquivos: Optional[list[dict]]) -> dict[str, dict]:
    saida: dict[str, dict] = {}
    for item in meta_arquivos or []:
        nome = item.get("arquivo") or item.get("nome")
        if not nome:
            continue
        chars = item.get("chars")
        try:
            chars_int = int(chars) if chars not in (None, "") else None
        except Exception:
            chars_int = None
        saida[nome] = {
            "tipo_extraido": item.get("tipo"),
            "chars_extraidos": chars_int,
        }
    return saida


def _persistir_arquivos_analise(analise_id: str, arquivos_raw: list[tuple[str, bytes]], meta_arquivos: Optional[list[dict]] = None) -> list[dict]:
    anexos: list[dict] = []
    meta_por_nome = _normalizar_meta_extracao(meta_arquivos)
    if not arquivos_raw:
        return anexos

    if _DATABASE_URL:
        try:
            with _db_conn() as conn:
                with conn.cursor() as cur:
                    cur.execute("DELETE FROM historico_arquivos WHERE analise_id = %s", (analise_id,))
                    for ordem, (nome, conteudo) in enumerate(arquivos_raw):
                        extra = meta_por_nome.get(nome, {})
                        meta = _normalizar_meta_arquivo(nome, conteudo, ordem, extra=extra)
                        cur.execute("""
                            INSERT INTO historico_arquivos (
                                id, analise_id, ordem, nome_original, mime_type,
                                tamanho_bytes, sha256, tipo_extraido, chars_extraidos, conteudo
                            )
                            VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                            ON CONFLICT (id) DO UPDATE SET
                                analise_id = EXCLUDED.analise_id,
                                ordem = EXCLUDED.ordem,
                                nome_original = EXCLUDED.nome_original,
                                mime_type = EXCLUDED.mime_type,
                                tamanho_bytes = EXCLUDED.tamanho_bytes,
                                sha256 = EXCLUDED.sha256,
                                tipo_extraido = EXCLUDED.tipo_extraido,
                                chars_extraidos = EXCLUDED.chars_extraidos,
                                conteudo = EXCLUDED.conteudo
                        """, (
                            meta["id"],
                            analise_id,
                            ordem,
                            nome,
                            meta["mime_type"],
                            meta["tamanho_bytes"],
                            meta["sha256"],
                            meta.get("tipo_extraido"),
                            meta.get("chars_extraidos"),
                            conteudo,
                        ))
                        anexos.append(meta)
                conn.commit()
            return anexos
        except Exception:
            logger.exception("Erro ao salvar arquivos do hist?rico", extra={"request_id": "-"})

    for ordem, (nome, conteudo) in enumerate(arquivos_raw):
        extra = meta_por_nome.get(nome, {})
        anexos.append(_normalizar_meta_arquivo(nome, conteudo, ordem, extra=extra))
    return anexos


def registrar_analise(ficha: str, arquivos_raw: Optional[list[tuple[str, bytes]]] = None, meta_arquivos: Optional[list[dict]] = None, fonte: str = "texto"):
    analise_id = uuid.uuid4().hex[:10]
    anexos = _persistir_arquivos_analise(analise_id, arquivos_raw or [], meta_arquivos)
    registro = {
        "id":        analise_id,
        "timestamp": datetime.datetime.now().isoformat(timespec="seconds"),
        "orgao":     extrair_campo(ficha, "Ãƒâ€œrgÃƒÂ£o"),
        "valor":     extrair_campo(ficha, "Valor Estimado Total"),
        "objeto":    extrair_objeto(ficha),
        "segmento":  detectar_segmento(ficha),
        "score":     extrair_score(ficha),
        "ficha":     ficha,
        "arquivos":  anexos,
        "fonte":     fonte,
    }
    _historico.insert(0, registro)
    if len(_historico) > 500:
        _historico.pop()
    _salvar_historico()


MAX_CHARS = 400_000

# Textos pequenos: Groq e OpenRouter primeiro (gratuitos), OpenAI sÃ³ se necessÃ¡rio
# Textos grandes: OpenAI primeiro (melhor qualidade para docs longos)
LIMITE_PEQUENO = 15_000  # chars â€” abaixo disso evita cobrar do OpenAI

PROVEDORES_PEQUENO = [
    ("gemini", "gemini-2.5-flash-lite",   400_000),
    ("gemini", "gemini-2.5-flash",        400_000),
    (_groq,    "llama-3.3-70b-versatile",  32_000),
]

PROVEDORES_GRANDE = [
    ("gemini", "gemini-2.5-flash-lite",   400_000),
    ("gemini", "gemini-2.5-flash",        400_000),
    (_groq,    "llama-3.3-70b-versatile",  32_000),
]


async def _chamar_gemini_http(texto: str, num_docs: int, modelo: str) -> str:
    if not _gemini_api_key:
        raise HTTPException(503, "GEMINI_API_KEY nao configurada.")

    texto_truncado = texto[:400_000]
    user_prompt = USER_TEMPLATE.format(num_docs=num_docs, texto=texto_truncado)
    payload = {
        "contents": [
            {
                "role": "user",
                "parts": [{"text": f"{SYSTEM_PROMPT}\n\n{user_prompt}"}],
            }
        ],
        "generationConfig": {
            "temperature": 0.2,
            "maxOutputTokens": 8192,
        },
    }
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{modelo}:generateContent"
    req = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Content-Type": "application/json",
            "x-goog-api-key": _gemini_api_key,
        },
        method="POST",
    )

    def _do_request():
        with urllib.request.urlopen(req, timeout=90) as resp:
            return resp.read().decode("utf-8")

    try:
        raw = await asyncio.to_thread(_do_request)
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="ignore") if hasattr(e, "read") else ""
        raise HTTPException(e.code, f"Erro ao chamar Gemini: {body or str(e)}")
    except urllib.error.URLError as e:
        raise HTTPException(500, f"Erro de conexao com Gemini: {e.reason}")

    data = json.loads(raw)
    candidatos = data.get("candidates") or []
    if not candidatos:
        raise HTTPException(503, "Gemini retornou resposta vazia.")
    parts = (
        candidatos[0]
        .get("content", {})
        .get("parts", [])
    )
    ficha = "".join(part.get("text", "") for part in parts).strip()
    if not ficha:
        raise HTTPException(503, "Gemini retornou texto vazio.")
    idx = ficha.find("## FICHA")
    if idx > 0:
        ficha = ficha[idx:]
    return ficha


async def chamar_groq(texto: str, num_docs: int) -> str:
    ultimo_erro = ""
    provedores = PROVEDORES_PEQUENO if len(texto) <= LIMITE_PEQUENO else PROVEDORES_GRANDE

    for cliente, modelo, max_chars in provedores:
        if cliente is None:
            ultimo_erro = "provedor sem API key configurada"
            continue
        if cliente == "gemini":
            try:
                ficha = await _chamar_gemini_http(texto[:max_chars], num_docs, modelo)
                if ficha.startswith("## FICHA"):
                    _stats["total_analises"] += 1
                    p = _stats["por_provedor"].setdefault(
                        modelo,
                        {"analises": 0, "tokens": 0, "custo_usd": 0.0},
                    )
                    p["analises"] += 1
                    return ficha
                ultimo_erro = f"Modelo {modelo} nao seguiu o formato esperado"
                continue
            except HTTPException as e:
                if e.status_code in (400, 401, 402, 404, 413, 429, 503):
                    ultimo_erro = "alguns minutos"
                    continue
                raise

        texto_truncado = texto[:max_chars]
        user_prompt = USER_TEMPLATE.format(num_docs=num_docs, texto=texto_truncado)
        try:
            response = await cliente.chat.completions.create(
                model=modelo,
                max_tokens=8192,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": user_prompt},
                ],
            )

            if not response.choices:
                ultimo_erro = f"Modelo {modelo} retornou resposta vazia"
                continue
            ficha = (response.choices[0].message.content or "").strip()
            idx = ficha.find("## FICHA")
            if idx > 0:
                ficha = ficha[idx:]
            if ficha.startswith("## FICHA"):
                _stats["total_analises"] += 1
                uso = getattr(response, "usage", None)
                tok_in  = getattr(uso, "prompt_tokens", 0) or 0
                tok_out = getattr(uso, "completion_tokens", 0) or 0
                ci, co  = _CUSTO.get(modelo, (0.0, 0.0))
                custo   = (tok_in * ci + tok_out * co) / 1_000_000
                _stats["tokens_input_total"]  += tok_in
                _stats["tokens_output_total"] += tok_out
                _stats["custo_usd_total"]     += custo
                p = _stats["por_provedor"].setdefault(modelo, {"analises": 0, "tokens": 0, "custo_usd": 0.0})
                p["analises"] += 1
                p["tokens"]   += tok_in + tok_out
                p["custo_usd"] = round(p["custo_usd"] + custo, 6)
                return ficha
            ultimo_erro = f"Modelo {modelo} nÃ£o seguiu o formato esperado"
            continue
        except openai.APITimeoutError:
            raise HTTPException(504, "Tempo limite excedido ao chamar a API.")
        except openai.APIConnectionError:
            raise HTTPException(500, "Erro de conexÃ£o com a API.")
        except openai.APIStatusError as e:
            if e.status_code in (400, 401, 402, 404, 413, 429, 503):
                msg = str(e.message)
                # extrai tempo de retry da mensagem da API
                m = re.search(r'retry in (\d+(?:\.\d+)?)\s*s', msg, re.IGNORECASE)
                if m:
                    segundos = int(float(m.group(1))) + 1
                    if segundos >= 3600:
                        tempo = f"{segundos // 3600}h"
                    elif segundos >= 60:
                        tempo = f"{segundos // 60} minutos"
                    else:
                        tempo = f"{segundos} segundos"
                    ultimo_erro = tempo
                elif 'per-day' in msg or 'per_day' in msg or 'daily' in msg.lower():
                    ultimo_erro = "algumas horas (limite diÃ¡rio atingido)"
                else:
                    ultimo_erro = "alguns minutos"
                continue
            raise HTTPException(500, f"Erro ao chamar a API: {e.message}")

    raise HTTPException(503, f"Todas as IAs estÃ£o sobrecarregadas no momento. Tente novamente em {ultimo_erro or 'alguns minutos'}.")


def _registrar_uso_parser_local(confianca: int):
    _stats["total_analises"] += 1
    p = _stats["por_provedor"].setdefault(
        "parser-local",
        {"analises": 0, "tokens": 0, "custo_usd": 0.0},
    )
    p["analises"] += 1
    p["tokens"] = 0
    p["custo_usd"] = 0.0
    p["confianca"] = confianca


def _http_post_json(url: str, payload: dict, timeout_s: float) -> dict:
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    req = urllib.request.Request(
        url,
        data=data,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=timeout_s) as resp:
        raw = resp.read()
    try:
        return json.loads(raw.decode("utf-8", errors="replace"))
    except Exception:
        return {"_raw": raw.decode("utf-8", errors="replace")}


async def _chamar_llm_local_chat(messages: list[dict], model: str) -> str:
    if not LOCAL_LLM_ENABLED:
        raise HTTPException(503, "LLM local desativado.")

    if LOCAL_LLM_PROVIDER == "ollama":
        url = LOCAL_LLM_BASE_URL.rstrip("/") + "/api/chat"
        payload = {"model": model, "messages": messages, "stream": False}
        try:
            data = await asyncio.to_thread(_http_post_json, url, payload, LOCAL_LLM_TIMEOUT_S)
        except Exception as e:
            raise HTTPException(503, f"Falha ao chamar LLM local (Ollama): {e}")
        content = (
            (data.get("message") or {}).get("content")
            or (data.get("response") or "")
            or ""
        )
        return str(content).strip()

    if LOCAL_LLM_PROVIDER == "openai_compat":
        url = LOCAL_LLM_BASE_URL.rstrip("/") + "/v1/chat/completions"
        payload = {"model": model, "messages": messages, "temperature": 0.2, "stream": False}
        try:
            data = await asyncio.to_thread(_http_post_json, url, payload, LOCAL_LLM_TIMEOUT_S)
        except Exception as e:
            raise HTTPException(503, f"Falha ao chamar LLM local (OpenAI-compat): {e}")
        choices = data.get("choices") or []
        if choices:
            content = (((choices[0] or {}).get("message") or {}).get("content") or "").strip()
            return content
        return ""

    raise HTTPException(500, f"LOCAL_LLM_PROVIDER inválido: {LOCAL_LLM_PROVIDER}")


def _montar_ficha_a_partir_de_campos(campos: dict) -> str:
    # Mantém compatibilidade com o restante do app (render e histórico).
    return gerar_ficha(campos)


async def tentar_fallback_llm_local(texto: str, num_docs: int) -> str | None:
    """
    Fallback local para extrair campos principais quando o parser falhar.
    Retorna uma ficha em Markdown ou None.
    """
    if not LOCAL_LLM_ENABLED:
        return None

    system = (
        "Você é um extrator de informações de editais brasileiros.\n"
        "Tarefa: extrair os campos principais de um edital (podem existir anexos misturados).\n"
        "\n"
        "Regras:\n"
        "1) Responda APENAS em JSON válido (sem markdown).\n"
        "2) Para cada campo, inclua também uma evidência curta (trecho literal do texto) que justifique.\n"
        "3) Se houver conflito (valores unitários vs total), priorize o TOTAL do contrato/contratação.\n"
        "4) Para datas, priorize a data/hora da sessão pública/abertura de propostas.\n"
        "\n"
        "Formato de saída:\n"
        "{\n"
        '  \"numero_edital\": \"...\",\n'
        '  \"orgao\": \"...\",\n'
        '  \"modalidade\": \"...\",\n'
        '  \"objeto\": \"...\",\n'
        '  \"valor\": \"...\",\n'
        '  \"data_abertura\": \"...\",\n'
        '  \"criterio_julgamento\": \"...\",\n'
        '  \"prazo_vigencia\": \"...\",\n'
        '  \"evidencias\": {\n'
        '    \"numero_edital\": \"...\",\n'
        '    \"orgao\": \"...\",\n'
        '    \"valor\": \"...\",\n'
        '    \"data_abertura\": \"...\"\n'
        "  }\n"
        "}\n"
        "\n"
        "Use 'Não identificado' quando não encontrar."
    )
    user = (
        f"Documentos no caso: {num_docs}\n\n"
        "Texto consolidado do edital (pode conter anexos):\n"
        + texto[:120000]
    )
    content = await _chamar_llm_local_chat(
        [{"role": "system", "content": system}, {"role": "user", "content": user}],
        LOCAL_LLM_MODEL,
    )
    if not content:
        return None
    try:
        data = json.loads(content)
    except Exception:
        # tenta recuperar JSON dentro de texto maior
        m = re.search(r"\{.*\}", content, re.DOTALL)
        if not m:
            return None
        try:
            data = json.loads(m.group(0))
        except Exception:
            return None

    evid = data.get("evidencias") or {}

    def _pick_best_valor(texto_all: str, valor_llm: str) -> str:
        # tenta puxar o valor total quando houver marcador forte no texto
        t = texto_all
        m = re.search(r"(?:valor\\s+total\\s+da\\s+contrata[cç][aã]o|total\\s+da\\s+contrata[cç][aã]o|valor\\s+estimado\\s+total)[^\\n]{0,80}?(R\\$\\s*[\\d\\.]{1,3}(?:\\.[\\d]{3})*(?:,\\d{2})?)", t, re.IGNORECASE)
        if m:
            return m.group(1).strip()
        return valor_llm

    def _pick_best_data(texto_all: str, data_llm: str) -> str:
        t = texto_all
        m = re.search(r"(?:data\\s+da\\s+sess[aã]o\\s+p[uú]blica|abertura\\s+da\\s+sess[aã]o|sess[aã]o\\s+p[uú]blica)[^\\n]{0,80}?(\\d{2}/\\d{2}/\\d{4})", t, re.IGNORECASE)
        if m:
            return m.group(1).strip()
        return data_llm

    valor_llm = (data.get("valor") or "").strip() or "Não identificado"
    data_llm = (data.get("data_abertura") or "").strip() or "Não identificado"
    valor_corr = _pick_best_valor(texto, valor_llm)
    data_corr = _pick_best_data(texto, data_llm)

    campos = {
        "numero_edital": (data.get("numero_edital") or "").strip() or "Não identificado",
        "orgao": (data.get("orgao") or "").strip() or "Não identificado",
        "cnpj": "Não identificado",
        "modalidade": (data.get("modalidade") or "").strip() or "Não identificado",
        "objeto": (data.get("objeto") or "").strip() or "Não identificado",
        "valor": valor_corr,
        "data_abertura": data_corr,
        "prazo_envio_proposta": "Não identificado",
        "prazo_vigencia": (data.get("prazo_vigencia") or "").strip() or "Não identificado",
        "criterio_julgamento": (data.get("criterio_julgamento") or "").strip() or "Não identificado",
        "documentos_habilitacao": [],
        "segmento": detectar_segmento((data.get("objeto") or "") + "\n" + texto),
    }
    confianca, faltantes = calcular_confianca(campos)
    campos["confianca"] = confianca
    campos["faltantes"] = faltantes
    campos["usar_fallback_api"] = False
    score, nivel, justificativas = calcular_score_viabilidade(campos)
    campos["score"] = score
    campos["nivel"] = nivel
    campos["justificativas_score"] = justificativas
    return _montar_ficha_a_partir_de_campos(campos)


async def _enriquecer_cnpj(cnpj: str) -> dict | None:
    """Consulta BrasilAPI para obter razÃ£o social oficial do CNPJ. Gratuito, sem auth."""
    cnpj_limpo = re.sub(r"\D", "", cnpj)
    if len(cnpj_limpo) != 14:
        return None
    url = f"https://brasilapi.com.br/api/cnpj/v1/{cnpj_limpo}"
    try:
        loop = asyncio.get_event_loop()
        def _fetch():
            req = urllib.request.Request(url, headers={"User-Agent": "LicitaPRO/1.0"})
            with urllib.request.urlopen(req, timeout=5) as resp:
                return json.loads(resp.read().decode())
        return await loop.run_in_executor(None, _fetch)
    except Exception as e:
        logger.warning("Falha ao consultar CNPJ %s: %s", cnpj_limpo, e, extra={"request_id": "-"})
        return None


async def analisar_com_fallback(texto: str, num_docs: int, modo: str = "auto") -> str:
    # modo "ia": pula o parser e vai direto para IA
    if modo == "ia" or not USAR_PARSER_LOCAL:
        logger.info("Analise por IA (modo=%s)", modo, extra={"request_id": "-"})
        return await chamar_groq(texto, num_docs)

    try:
        resultado = analisar_sem_api(texto, min_confianca=PARSER_MIN_CONFIANCA)
    except Exception:
        logger.exception("Erro no parser local", extra={"request_id": "-"})
        if modo != "parser" and PARSER_FALLBACK_API:
            return await chamar_groq(texto, num_docs)
        raise HTTPException(500, "Erro ao analisar edital pelo parser local.")

    texto_longo = len(texto) > PARSER_MAX_CHARS_FALLBACK

    # modo "parser": retorna o resultado do parser sem nunca chamar IA
    if modo == "parser":
        logger.info(
            "Analise por parser (confiança=%s%%)",
            resultado.get("confianca", 0),
            extra={"request_id": "-"},
        )
    else:
        # modo "auto": fallback para IA se confiança baixa
        if resultado.get("usar_fallback_api") and PARSER_FALLBACK_API:
            # 1) tenta fallback local (se habilitado) antes de gastar API externa
            if LOCAL_LLM_ENABLED:
                try:
                    ficha_local = await tentar_fallback_llm_local(texto, num_docs)
                    if ficha_local and ficha_local.startswith("## FICHA"):
                        logger.info("LLM local usado (sem API externa).", extra={"request_id": "-"})
                        _stats["total_analises"] += 1
                        p = _stats["por_provedor"].setdefault(
                            "llm-local",
                            {"analises": 0, "tokens": 0, "custo_usd": 0.0},
                        )
                        p["analises"] += 1
                        return ficha_local
                except HTTPException as e:
                    logger.info("LLM local indisponivel: %s", e.detail, extra={"request_id": "-"})
                except Exception as e:
                    logger.warning("Erro no LLM local: %s", e, extra={"request_id": "-"})

            # 2) fallback API (com truncamento quando necessário)
            if texto_longo:
                corte = int(PARSER_MAX_CHARS_FALLBACK * 0.60)
                texto_api = (
                    texto[:corte]
                    + "\n\n[...trecho omitido...]\n\n"
                    + texto[-(PARSER_MAX_CHARS_FALLBACK - corte):]
                )
                logger.info(
                    "Confiança baixa (%s%%), doc longo (%s chars) — truncado para fallback API.",
                    resultado.get("confianca", 0),
                    len(texto),
                    extra={"request_id": "-"},
                )
            else:
                texto_api = texto
                logger.info(
                    "Confiança baixa (%s%%). Usando fallback por API.",
                    resultado.get("confianca", 0),
                    extra={"request_id": "-"},
                )
            return await chamar_groq(texto_api, num_docs)

    if resultado.get("usar_fallback_api") and texto_longo and (modo == "parser" or not PARSER_FALLBACK_API):
        logger.info(
            "Documento longo demais para fallback automático; mantendo saída do parser local.",
            extra={"request_id": "-"},
        )
    # enriquecimento: BrasilAPI CNPJ â†’ razÃ£o social oficial
    cnpj_extraido = resultado.get("cnpj", "")
    if _is_identificado(cnpj_extraido):
        dados_cnpj = await _enriquecer_cnpj(cnpj_extraido)
        if dados_cnpj:
            razao = (dados_cnpj.get("razao_social") or "").strip()
            orgao_atual = resultado.get("orgao", "")
            # usa razÃ£o social oficial quando: (a) orgÃ£o nÃ£o identificado,
            # (b) orgÃ£o tem mais de 10 palavras (capturou contexto demais),
            # (c) orgÃ£o contÃ©m artefatos de tabela PDF ("PROJETO:", "INFORMAÃ‡Ã•ES")
            orgao_ruim = (
                not _is_identificado(orgao_atual)
                or len(orgao_atual.split()) > 10
                or any(s in orgao_atual.upper() for s in ("PROJETO:", "INFORMAÃ‡Ã•ES", "ORIENTADA"))
            )
            if razao and orgao_ruim:
                resultado["orgao"] = razao.title()
                resultado["ficha"] = gerar_ficha(resultado)
                logger.info("Ã“rgÃ£o enriquecido via CNPJ: %s", razao[:80], extra={"request_id": "-"})
    _registrar_uso_parser_local(int(resultado.get("confianca", 0)))
    return resultado["ficha"]


@app.get("/status", response_class=HTMLResponse)
async def status():
    total   = _stats["total_analises"]
    hoje    = _stats["analises_hoje"]
    hist_n  = len(_historico)
    deploy_label = APP_DEPLOYED_AT.strip() if APP_DEPLOYED_AT else "nÃ£o informado"
    data_reset = _stats["hoje"]

    return f"""<!DOCTYPE html>
<html lang="pt-BR">
<head>
<meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>Status â€” LicitaPro</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap" rel="stylesheet">
<style>
:root{{
  --brand-900:#061A33;--brand-800:#0A2540;--brand-600:#1E4FA0;--brand-500:#2F6FE0;
  --brand-100:#E6EEFC;--brand-50:#F4F8FE;
  --ink-900:#111827;--ink-800:#1F2937;--ink-600:#4B5563;--ink-500:#6B7280;
  --ink-400:#9CA3AF;--ink-300:#D1D5DB;--ink-200:#E5E7EB;--ink-150:#EEF0F3;
  --ink-100:#F3F4F6;--ink-50:#F9FAFB;--ink-0:#FFFFFF;
  --success-700:#166534;--success-100:#DCFCE7;
  --warn-700:#92400E;--warn-500:#F59E0B;--warn-50:#FFFBEB;
  --danger-700:#991B1B;--danger-500:#EF4444;--danger-50:#FEF2F2;
  --bg:#FFFFFF;--bg-subtle:#FAFBFC;--surface:var(--ink-0);
  --fg-1:var(--ink-900);--fg-2:var(--ink-600);--fg-3:var(--ink-500);--fg-4:var(--ink-400);
  --border:var(--ink-200);--border-subtle:var(--ink-150);
  --font-sans:'Inter',-apple-system,BlinkMacSystemFont,'Segoe UI',sans-serif;
  --font-mono:'JetBrains Mono',Menlo,Consolas,monospace;
  --radius-sm:6px;--radius-md:10px;--radius-lg:14px;--radius-full:9999px;
  --shadow-sm:0 1px 2px rgba(11,15,20,.04),0 1px 3px rgba(11,15,20,.06);
  --ease-out:cubic-bezier(0.22,1,0.36,1);
}}
*{{margin:0;padding:0;box-sizing:border-box}}
body{{font-family:var(--font-sans);background:var(--bg-subtle);color:var(--fg-1);min-height:100vh;-webkit-font-smoothing:antialiased}}

nav{{position:sticky;top:0;z-index:10;background:rgba(255,255,255,.72);backdrop-filter:blur(20px) saturate(180%);-webkit-backdrop-filter:blur(20px) saturate(180%);border-bottom:1px solid var(--border-subtle)}}
.nav-inner{{max-width:960px;margin:0 auto;padding:0 24px;height:56px;display:flex;align-items:center;justify-content:space-between;gap:16px}}
.logo{{font-size:17px;font-weight:700;color:var(--ink-900);letter-spacing:-.025em}}
.logo-dot{{color:var(--brand-500)}}
.btn-back{{display:inline-flex;align-items:center;gap:6px;padding:7px 13px;border-radius:var(--radius-md);border:1px solid var(--border);background:transparent;color:var(--fg-2);font-size:13px;font-weight:500;font-family:var(--font-sans);cursor:pointer;text-decoration:none;transition:background .12s}}
.btn-back:hover{{background:var(--ink-50);color:var(--fg-1)}}

main{{max-width:960px;margin:0 auto;padding:32px 24px 64px}}
.page-header{{margin-bottom:28px}}
.page-title{{font-size:22px;font-weight:600;color:var(--ink-900);letter-spacing:-.025em}}
.page-sub{{font-size:13px;color:var(--fg-3);margin-top:4px;font-variant-numeric:tabular-nums}}

.grid{{display:grid;gap:14px;margin-bottom:14px}}
.g4{{grid-template-columns:repeat(4,1fr)}}
@media(max-width:720px){{.g4{{grid-template-columns:1fr 1fr}}}}
@media(max-width:460px){{.g4{{grid-template-columns:1fr}}}}

.card{{background:var(--surface);border:1px solid var(--border);border-radius:var(--radius-lg);padding:20px 22px;box-shadow:var(--shadow-sm)}}
.stat-label{{font-size:11px;font-weight:600;letter-spacing:.06em;text-transform:uppercase;color:var(--fg-3);margin-bottom:8px}}
.stat-val{{font-size:32px;font-weight:700;color:var(--ink-900);line-height:1;letter-spacing:-.025em;font-variant-numeric:tabular-nums}}
.stat-val-sm{{font-size:24px;font-weight:700;color:var(--ink-900);line-height:1;letter-spacing:-.015em;font-variant-numeric:tabular-nums}}
.stat-sub{{font-size:12px;color:var(--fg-4);margin-top:5px;font-variant-numeric:tabular-nums}}

.footer-note{{text-align:right;font-size:12px;color:var(--fg-4);margin-top:12px;font-variant-numeric:tabular-nums}}
</style>
</head>
<body>
<nav>
  <div class="nav-inner">
    <span class="logo">Licita<span class="logo-dot">Pro</span></span>
    <a class="btn-back" href="/">
      <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round"><path d="m15 18-6-6 6-6"/></svg>
      Voltar
    </a>
  </div>
</nav>
<main>
  <div class="page-header">
    <h1 class="page-title">Status do sistema</h1>
    <p class="page-sub">SessÃ£o atual &nbsp;Â·&nbsp; Reseta em: {data_reset}</p>
  </div>

  <div class="grid g4">
    <div class="card">
      <div class="stat-label">AnÃ¡lises hoje</div>
      <div class="stat-val">{hoje}</div>
      <div class="stat-sub">anÃ¡lises processadas hoje</div>
    </div>
    <div class="card">
      <div class="stat-label">Total na sessÃ£o</div>
      <div class="stat-val">{total}</div>
      <div class="stat-sub">anÃ¡lises realizadas</div>
    </div>
    <div class="card">
      <div class="stat-label">HistÃ³rico salvo</div>
      <div class="stat-val">{hist_n}</div>
      <div class="stat-sub">anÃ¡lises no arquivo</div>
    </div>
    <div class="card">
      <div class="stat-label">VersÃ£o vigente</div>
      <div class="stat-val-sm">{APP_VERSION_LABEL}</div>
      <div class="stat-sub">commit {APP_COMMIT_LABEL} Â· deploy {deploy_label}</div>
    </div>
  </div>

  <p class="footer-note">Atualiza automaticamente a cada 30 s &nbsp;Â·&nbsp; {_stats["hoje"]} &nbsp;Â·&nbsp; {APP_VERSION_LABEL} &nbsp;Â·&nbsp; commit {APP_COMMIT_LABEL}</p>
</main>
<script>setTimeout(()=>location.reload(),30000);</script>
</body>
</html>"""


@app.post("/importar/arquivo")
async def importar_arquivo(request: Request, arquivos: list[UploadFile] = File(...)):
    if not arquivos:
        raise HTTPException(400, "Nenhum arquivo enviado.")
    _audit("importar_arquivo_start", request.state.request_id, arquivos=len(arquivos))
    importados, ignorados = [], []
    for arq in arquivos:
        conteudo = await arq.read()
        try:
            texto = extrair_texto(arq.filename, conteudo)
        except Exception as e:
            _audit("importar_arquivo_read_error", request.state.request_id, arquivo=arq.filename, erro=str(e))
            ignorados.append({"arquivo": arq.filename, "motivo": "erro ao ler: " + str(e)})
            continue
        texto = (texto or "").strip()
        if not texto:
            ignorados.append({"arquivo": arq.filename, "motivo": "nenhum texto extraído do arquivo"})
            continue
        if "## FICHA" in texto:
            ficha = texto[texto.find("## FICHA"):]
        else:
            ficha = "## FICHA DE LICITAÇÃO\n\n" + texto
        registrar_analise(ficha, arquivos_raw=[(arq.filename, conteudo)], fonte="importacao")
        importados.append(arq.filename)
    _audit("importar_arquivo_done", request.state.request_id, importados=len(importados), ignorados=len(ignorados))
    return {"importados": importados, "ignorados": ignorados}


@app.post("/importar/texto")
async def importar_texto(http_request: Request, request: AnalisarRequest):
    texto = request.texto.strip()
    if not texto:
        raise HTTPException(400, "Texto vazio.")
    _audit("importar_texto_start", http_request.state.request_id, chars=len(texto), num_docs=request.num_docs)
    ficha = texto[texto.find("## FICHA"):] if "## FICHA" in texto else "## FICHA DE LICITAÇÃO\n\n" + texto
    registrar_analise(ficha)
    _audit("importar_texto_done", http_request.state.request_id, chars=len(texto), num_docs=request.num_docs)
    return {"ok": True}


@app.post("/api/reclassificar")
async def api_reclassificar():
    antes = {r["id"]: r.get("segmento") for r in _historico}
    _reclassificar_historico()
    atualizados = sum(1 for r in _historico if antes.get(r["id"]) != r.get("segmento"))
    return {"atualizados": atualizados, "total": len(_historico)}


@app.get("/historico")
async def get_historico():
    return {"historico": [{k: v for k, v in r.items() if k != "ficha"} for r in _historico]}

@app.get("/historico/{id}")
async def get_ficha_historico(id: str):
    for r in _historico:
        if r["id"] == id:
            return {"ficha": r.get("ficha"), "orgao": r.get("orgao"), "segmento": r.get("segmento"), "score": r.get("score"), "arquivos": r.get("arquivos", []), "fonte": r.get("fonte")}
    raise HTTPException(404, "Análise não encontrada.")


@app.get("/historico/{id}/arquivos/{arquivo_id}")
async def baixar_arquivo_historico(id: str, arquivo_id: str):
    if _DATABASE_URL:
        try:
            with _db_conn() as conn:
                with conn.cursor() as cur:
                    cur.execute("""
                        SELECT nome_original, mime_type, conteudo
                        FROM historico_arquivos
                        WHERE analise_id = %s AND id = %s
                    """, (id, arquivo_id))
                    row = cur.fetchone()
                    if row:
                        nome, mime_type, conteudo = row
                        return Response(
                            content=bytes(conteudo),
                            media_type=mime_type or "application/octet-stream",
                            headers={"Content-Disposition": f'attachment; filename="{nome}"'}
                        )
        except Exception:
            logger.exception("Erro ao baixar arquivo do hist?rico", extra={"request_id": "-"})
    raise HTTPException(404, "Arquivo nÃ£o encontrado.")


@app.get("/", response_class=HTMLResponse)
async def root():
    return (
        HTML_PAGE
        .replace("{APP_VERSION_LABEL}", APP_VERSION_LABEL)
        .replace("{APP_COMMIT_LABEL}", APP_COMMIT_LABEL)
    )


@app.post("/analisar/arquivo", response_model=AnalisarResponse)
async def analisar_arquivo(request: Request, arquivos: list[UploadFile] = File(...), modo: str = Form("auto")):
    if not arquivos:
        raise HTTPException(400, "Nenhum arquivo enviado.")
    _audit("analisar_arquivo_start", request.state.request_id, arquivos=len(arquivos))

    hoje = datetime.date.today().isoformat()
    if _stats["hoje"] != hoje:
        _stats["hoje"] = hoje
        _stats["analises_hoje"] = 0

    if _stats["analises_hoje"] >= LIMITE_DIARIO:
        raise HTTPException(429, f"Limite diário de {LIMITE_DIARIO} análises atingido. Tente novamente amanhã.")

    arquivos_raw: list[tuple[str, bytes]] = []
    for arq in arquivos:
        try:
            conteudo = await arq.read()
        except Exception as e:
            _audit("analisar_arquivo_read_error", request.state.request_id, arquivo=arq.filename, erro=str(e))
            raise HTTPException(400, f"Erro ao ler '{arq.filename}': {e}")
        arquivos_raw.append((arq.filename, conteudo))

    try:
        texto_completo, meta_arquivos = montar_texto_caso_classificado_raw(arquivos_raw)
    except Exception as e:
        _audit("analisar_arquivo_process_error", request.state.request_id, erro=str(e))
        raise HTTPException(400, f"Erro ao processar arquivos enviados: {e}")

    if len(texto_completo) > MAX_CHARS:
        cota = MAX_CHARS // max(1, len(meta_arquivos))
        partes = []
        for nome, conteudo in arquivos_raw:
            try:
                _, texto = extrair_texto_com_tipo(nome, conteudo)
            except Exception as e:
                raise HTTPException(400, f"Erro ao ler '{nome}': {e}")
            texto = texto or ""
            if len(texto) <= cota:
                trecho = texto
            else:
                inicio = int(cota * 0.60)
                fim = cota - inicio
                trecho = texto[:inicio] + "\n\n[...]\n\n" + texto[-fim:]
            partes.append(f"=== {nome} ===\n{trecho}")
        texto_completo = "\n\n".join(partes)
    ficha = await analisar_com_fallback(texto_completo, len(arquivos), modo=modo)
    _stats["analises_hoje"] += 1
    registrar_analise(ficha, arquivos_raw=arquivos_raw, meta_arquivos=meta_arquivos, fonte="upload")
    _audit("analisar_arquivo_done", request.state.request_id, arquivos=len(arquivos), chars=len(texto_completo))
    return AnalisarResponse(ficha=ficha)


@app.post("/analisar", response_model=AnalisarResponse)
async def analisar(http_request: Request, request: AnalisarRequest):
    _audit("analisar_texto_start", http_request.state.request_id, chars=len(request.texto or ""), num_docs=request.num_docs)
    ficha = await analisar_com_fallback(request.texto, request.num_docs, modo=request.modo)
    registrar_analise(ficha)
    _audit("analisar_texto_done", http_request.state.request_id, chars=len(request.texto or ""), num_docs=request.num_docs)
    return AnalisarResponse(ficha=ficha)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
